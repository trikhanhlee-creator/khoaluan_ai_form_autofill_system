#!/usr/bin/env python3
"""Debug parser to see which function is used and what labels it returns"""
import sys
sys.path.insert(0, '.')
from app.services.word_parser import WordParser
from docx import Document
import os

# Test with the user's actual file
file_path = 'uploads/1_1771926182.710894_thu.docx'

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("\nAvailable files:")
    for f in os.listdir('uploads'):
        if f.endswith('.docx'):
            print(f"  - {f}")
    exit(1)

print(f"=== DEBUGGING {file_path} ===\n")

# Check raw paragraphs
doc = Document(file_path)
print("Raw paragraphs:")
for idx, para in enumerate(doc.paragraphs):
    text = para.text
    stripped = text.strip()
    if stripped:
        print(f"  {idx}: original={repr(text)}")
        print(f"      stripped={repr(stripped)}")

print("\n" + "="*60 + "\n")

# Test each parser individually
parser = WordParser.__new__(WordParser)

# Temporarily initialize doc only
from docx import Document
parser.file_path = file_path
parser.doc = Document(file_path)
parser.fields = []

# Test parse_paragraphs
print("Testing parse_paragraphs():")
para_fields = parser.parse_paragraphs()
print(f"  Returned {len(para_fields)} fields:")
for field in para_fields:
    print(f"    - label={repr(field.label)}")
    print(f"      name={repr(field.name)}")

print("\n" + "="*60 + "\n")

# Test parse_tables
print("Testing parse_tables():")
table_fields = parser.parse_tables()
print(f"  Returned {len(table_fields)} fields:")
for field in table_fields:
    print(f"    - label={repr(field.label)}")
    print(f"      name={repr(field.name)}")

print("\n" + "="*60 + "\n")

# Test parse_all_text_content
print("Testing parse_all_text_content():")
content_fields = parser.parse_all_text_content()
print(f"  Returned {len(content_fields)} fields:")
for field in content_fields:
    print(f"    - label={repr(field.label)}")
    print(f"      name={repr(field.name)}")

print("\n" + "="*60 + "\n")

# Now test the actual parse() method
parser2 = WordParser(file_path)
all_fields = parser2.parse()
print("Using parser.parse():")
print(f"  Returned {len(all_fields)} fields:")
for field in all_fields:
    has_quotes = '"' in field.label
    status = "❌ HAS QUOTES" if has_quotes else "✅"
    print(f"    {status}: label={repr(field.label)}")
    print(f"           name={repr(field.name)}")
