#!/usr/bin/env python3
"""Check exact character codes in the Word file"""
import sys
sys.path.insert(0, '.')
from docx import Document

file_path = 'uploads/1_1771926182.710894_thu.docx'

doc = Document(file_path)

print("Checking paragraph text and character codes:\n")
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"Para {idx}: {repr(text)}")
        print(f"  Length: {len(text)}")
        print(f"  Char codes: {[ord(c) for c in text]}")
        print(f"  Last 3 chars: {repr(text[-3:])}")
        print(f"  Last 3 char codes: {[ord(c) for c in text[-3:]]}")
        print()

# Write to file
with open('char_codes.txt', 'w', encoding='utf-8') as f:
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            f.write(f"Para {idx}: {repr(text)}\n")
            f.write(f"  Char codes: {[ord(c) for c in text]}\n")
            f.write(f"  Last chars analysis:\n")
            for i, c in enumerate(text[-5:]):
                f.write(f"    [{-(5-i)}]: {repr(c)} (code {ord(c)})\n")
            f.write("\n")

print("Results written to char_codes.txt")
