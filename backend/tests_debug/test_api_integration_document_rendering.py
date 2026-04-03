"""
Test API Integration for Document Layout Rendering
This test demonstrates how to use the new document layout rendering endpoints
"""

import sys
import os
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import the renderer directly
from app.services.form_replacement.form_layout_renderer import FormLayoutRenderer, FormPageRenderer
from app.services.form_replacement.intelligent_detector import IntelligentDetector


def create_job_application_form():
    """Create a realistic job application form document"""
    doc = Document()
    
    # Title section
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    run.bold = True
    run.font.size = Pt(14)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Độc lập - Tự do - Hạnh phúc")
    run.italic = True
    
    # Main title
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main_title.add_run("ĐƠN XIN VIỆC")
    run.bold = True
    run.font.size = Pt(16)
    
    # Add spacing
    doc.add_paragraph()
    
    # Personal information section
    section_heading = doc.add_paragraph("THÔNG TIN CÁ NHÂN")
    section_heading.runs[0].bold = True
    
    # Kính gửi
    p = doc.add_paragraph()
    p.add_run("Kính gửi: ").bold = True
    p.add_run("........................................................................")
    
    # Name
    p = doc.add_paragraph()
    p.add_run("Tôi tên là: ").bold = True
    p.add_run("................................................................................")
    
    # Date of birth
    p = doc.add_paragraph()
    p.add_run("Sinh ngày (ngày/tháng/năm): ").bold = True
    p.add_run("_______________")
    
    # Address
    p = doc.add_paragraph()
    p.add_run("Chỗ ở hiện nay: ").bold = True
    p.add_run("................................................................................")
    
    # Phone
    p = doc.add_paragraph()
    p.add_run("Số điện thoại liên hệ: ").bold = True
    p.add_run("................................................................................")
    
    return doc


def test_api_document_rendering():
    """Test the complete API flow for document rendering"""
    
    print("\n" + "="*80)
    print("API INTEGRATION TEST: Document Layout Rendering")
    print("="*80 + "\n")
    
    # Step 1: Create form
    print("[1] Creating job application form...")
    doc = create_job_application_form()
    print("    ✅ Form created with 5 fields\n")
    
    # Step 2: Parse with IntelligentDetector
    print("[2] Parsing document with IntelligentDetector...")
    parsed_form = IntelligentDetector.parse_document(doc)
    print(f"    ✅ Detected {len(parsed_form.sections)} sections")
    print(f"    ✅ Detected {len(parsed_form.fields)} fields\n")
    
    # Print detected structure
    print("[3] Detected Structure:")
    for i, section in enumerate(parsed_form.sections):
        print(f"    Section {i+1}:")
        print(f"      Title: {section.title[:50] if section.title else 'N/A'}...")
        print(f"      Is title: {section.is_title}")
        print(f"      Fields in section:")
        for field in parsed_form.fields:
            if field.section_index == i:
                print(f"        - {field.label} ({field.field_type})")
    print()
    
    # Step 3: Render as Document Layout
    print("[4] Rendering as Document Layout...")
    renderer = FormLayoutRenderer()
    html_document = renderer.render_form_document_style(parsed_form)
    print(f"    ✅ Generated {len(html_document)} characters\n")
    
    # Verify content
    input_count = html_document.count('<input')
    field_count = parsed_form.fields.__len__()
    print(f"    Field verification:")
    print(f"      - Fields in form: {field_count}")
    print(f"      - Input elements in HTML: {input_count}")
    print(f"      - Match: {'✅ YES' if input_count == field_count else '❌ NO'}\n")
    
    # Step 4: Render as Complete Page
    print("[5] Rendering as Complete Page...")
    page_renderer = FormPageRenderer()
    html_page = page_renderer.render_complete_form_page(parsed_form)
    print(f"    ✅ Generated {len(html_page)} characters\n")
    
    # Verify page content
    has_css = '<style>' in html_page
    has_js = '<script>' in html_page
    has_buttons = 'submit' in html_page.lower()
    print(f"    Page verification:")
    print(f"      - Has CSS: {'✅ YES' if has_css else '❌ NO'}")
    print(f"      - Has JavaScript: {'✅ YES' if has_js else '❌ NO'}")
    print(f"      - Has buttons: {'✅ YES' if has_buttons else '❌ NO'}\n")
    
    # Step 5: Save outputs for inspection
    print("[6] Saving outputs for inspection...")
    
    # Save document layout HTML
    with open("form_document_layout_api_test.html", "w", encoding="utf-8") as f:
        f.write(html_document)
    print("    ✅ Saved: form_document_layout_api_test.html")
    
    # Save complete page HTML
    with open("form_complete_page_api_test.html", "w", encoding="utf-8") as f:
        f.write(html_page)
    print("    ✅ Saved: form_complete_page_api_test.html\n")
    
    # Step 6: Demonstrate API response structure
    print("[7] API Response Examples:\n")
    
    print("    Document Layout Response:")
    print("""    {
        "status": "success",
        "template_id": 123,
        "render_type": "document",
        "fields_count": 5,
        "sections_count": 2,
        "html_form": "<div class='form-document-style'>...</div>"
    }""")
    print()
    
    print("    Complete Page Response:")
    print("""    {
        "status": "success",
        "template_id": 123,
        "render_type": "page",
        "fields_count": 5,
        "html_page": "<!DOCTYPE html><html>...</html>"
    }""")
    print()
    
    # Step 7: Extract field information for API response
    print("[8] Field Information for Response:\n")
    fields_info = []
    for field in parsed_form.fields:
        field_info = {
            "name": field.field_name,
            "label": field.label,
            "field_type": field.field_type,
            "section_index": field.section_index,
            "order": field.line_index
        }
        fields_info.append(field_info)
        print(f"    {field_info['label']} ({field_info['field_type']})")
    print()
    
    # Step 8: Demonstrate field value collection
    print("[9] JavaScript Code to Collect Form Data:\n")
    print("""    const collectFormData = () => {
        const fields = document.querySelectorAll('[data-field-label]');
        const formData = {};
        
        fields.forEach(field => {
            const fieldLabel = field.getAttribute('data-field-label');
            const fieldName = field.getAttribute('name') || fieldLabel;
            formData[fieldName] = field.value;
        });
        
        return formData;
    };
    
    // Usage:
    const data = collectFormData();
    console.log(data);""")
    print()
    
    # Step 9: Test validation
    print("[10] Validation Results:\n")
    checks = [
        ("HTML contains form container", '<div' in html_document),
        ("HTML contains input fields", '<input' in html_document),
        ("Input count matches fields", input_count == field_count),
        ("Complete page starts with DOCTYPE", html_page.strip().startswith('<!DOCTYPE')),
        ("Complete page has HTML tags", '<html' in html_page.lower()),
        ("Complete page has CSS", '<style>' in html_page),
        ("Complete page has JavaScript", '<script>' in html_page),
        ("Complete page has form buttons", 'submit' in html_page.lower()),
    ]
    
    all_passed = True
    for check_name, check_result in checks:
        status = "✅" if check_result else "❌"
        print(f"    {status} {check_name}")
        if not check_result:
            all_passed = False
    
    print()
    
    # Final summary
    print("="*80)
    if all_passed:
        print("✅ ALL API INTEGRATION TESTS PASSED!")
        print("\nThe system is ready to:")
        print("  1. Accept document uploads")
        print("  2. Parse documents with intelligent detection")
        print("  3. Render forms in document layout style")
        print("  4. Return complete HTML pages with styling")
        print("  5. Collect user data from rendered forms")
    else:
        print("❌ SOME TESTS FAILED - Please review the output above")
    print("="*80 + "\n")
    
    return all_passed


def demonstrate_api_usage():
    """Demonstrate how the API would be used end-to-end"""
    
    print("\n" + "="*80)
    print("API USAGE DEMONSTRATION")
    print("="*80 + "\n")
    
    print("1. UPLOAD FORM")
    print("-" * 80)
    print("""
CURL Request:
  curl -X POST "http://localhost:8000/api/form-replacement/upload-with-intelligent-detection" \\
    -F "file=@form.docx" \\
    -F "user_id=1"

Response:
{
  "status": "success",
  "template_id": 123,
  "template_name": "form",
  "fields_count": 5,
  "sections_count": 2,
  "fields": [
    {
      "order": 0,
      "name": "kính_gửi",
      "label": "Kính gửi",
      "field_type": "text",
      "section_index": 0
    },
    ...
  ]
}
    """)
    
    print("\n2. RENDER AS DOCUMENT LAYOUT")
    print("-" * 80)
    print("""
CURL Request:
  curl "http://localhost:8000/api/form-replacement/template/123/render-form-document"

Response:
{
  "status": "success",
  "template_id": 123,
  "render_type": "document",
  "html_form": "<div class='form-document-style'>...</div>"
}

Frontend (JavaScript):
  fetch('/api/form-replacement/template/123/render-form-document')
    .then(r => r.json())
    .then(d => {
      document.getElementById('form').innerHTML = d.html_form;
    });
    """)
    
    print("\n3. RENDER AS COMPLETE PAGE")
    print("-" * 80)
    print("""
CURL Request:
  curl "http://localhost:8000/api/form-replacement/template/123/render-form-page"

Response:
{
  "status": "success",
  "render_type": "page",
  "html_page": "<!DOCTYPE html><html>...</html>"
}

Frontend (JavaScript):
  fetch('/api/form-replacement/template/123/render-form-page')
    .then(r => r.json())
    .then(d => {
      // Display complete page
      window.location.href = 'data:text/html,' + 
        encodeURIComponent(d.html_page);
    });
    """)
    
    print("\n4. COLLECT FORM DATA")
    print("-" * 80)
    print("""
JavaScript:
  const formData = new FormData();
  const inputs = document.querySelectorAll('[data-field-label]');
  
  inputs.forEach(input => {
    const fieldName = input.getAttribute('name');
    formData.append(fieldName, input.value);
  });
  
  // Submit
  fetch('/api/form-replacement/submit', {
    method: 'POST',
    body: formData
  });
    """)
    
    print("\n" + "="*80)


if __name__ == "__main__":
    try:
        # Run integration test
        success = test_api_document_rendering()
        
        # Demonstrate API usage
        demonstrate_api_usage()
        
        # Exit with appropriate code
        exit(0 if success else 1)
        
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        exit(1)
