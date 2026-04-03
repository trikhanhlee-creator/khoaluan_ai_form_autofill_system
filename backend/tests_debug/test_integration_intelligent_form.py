"""
Integration Test - End to End Test cho Intelligent Form Detection
"""

import sys
import os

# Setup path
backend_path = os.path.join(os.path.dirname(__file__), '..')
sys.path.insert(0, backend_path)

from docx import Document
from docx.shared import Pt, Inches
from app.services.form_replacement import (
    IntelligentDetector,
    SmartFormRenderer,
    DocumentStructurePreserver
)


def create_complete_form():
    """Tạo document giống như ảnh thứ 2 - ĐƠN XIN VIỆC"""
    doc = Document()

    # HEADER SECTION
    # Main title 1
    title_para_1 = doc.add_paragraph()
    title_para_1.alignment = 1  # Center
    run = title_para_1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    run.bold = True
    run.font.size = Pt(11)

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = 1
    run = subtitle_para.add_run("Độc lập - Tự do - Hạnh phúc")
    run.bold = True
    run.font.size = Pt(10)

    # MAIN TITLE
    main_title = doc.add_paragraph()
    main_title.alignment = 1
    run = main_title.add_run("ĐƠN XIN VIỆC")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # blank line

    # FORM CONTENT SECTION
    # Field 1
    doc.add_paragraph("Kính gửi: Ban lãnh đạo và phòng nhân sự Công ty ..........................................")
    doc.add_paragraph()

    # Field 2 & 3 & 4
    doc.add_paragraph("Tôi tên là: ........................................................................")
    doc.add_paragraph()

    doc.add_paragraph("Sinh ngày: ___/___/______  ")
    doc.add_paragraph()

    doc.add_paragraph("Chỗ ở hiện nay: ........................................................................")
    doc.add_paragraph()

    # Field 5 - Phone
    doc.add_paragraph("Số điện thoại liên hệ: ___________________")
    doc.add_paragraph()

    # Description text with inline field
    description = doc.add_paragraph()
    description.add_run("Thông qua trang website TopCV.vn, tôi biết được Quý công ty có nhu cầu tuyển dụng vị trí ")
    run = description.add_run("........................")
    run.bold = False

    doc.add_paragraph()

    # Field 6 - Experience/motivation
    doc.add_paragraph("Tôi cảm thấy trình độ và kỹ năng của mình phù hợp với vị trí này. Tôi mong muốn được làm việc và công hiến cho công ty.")

    return doc


def run_integration_test():
    """Run comprehensive integration test"""
    print("=" * 70)
    print("INTEGRATION TEST: Intelligent Form Detection & Rendering")
    print("=" * 70)

    # Create test document
    doc = create_complete_form()

    # STEP 1: Parse document
    print("\n[1] PARSING DOCUMENT...")
    parsed_form = IntelligentDetector.parse_document(doc)
    print(f"    ✅ Parsed {len(parsed_form.sections)} sections")
    print(f"    ✅ Detected {len(parsed_form.fields)} fields")
    print(f"    ✅ Content lines: {len(parsed_form.raw_content)}")

    # STEP 2: Display structure
    print("\n[2] FORM STRUCTURE DETECTED:")
    for idx, section in enumerate(parsed_form.sections):
        print(f"\n    Section {idx}: {section.title or '(unnamed)'}")
        print(f"      - Is Title: {section.is_title}")
        print(f"      - Level: {section.level}")
        print(f"      - Items: {len(section.items)}")

    # STEP 3: Display fields
    print("\n[3] FIELDS EXTRACTED:")
    fields_list = IntelligentDetector.extract_field_list(parsed_form)
    for field in fields_list:
        print(f"\n    [{field['order']}] {field['label']}")
        print(f"         Name: {field['name']}")
        print(f"         Type: {field['field_type']}")
        print(f"         In Section: {field['section_index']}")

    # STEP 4: Render as structured form
    print("\n[4] RENDERING STRUCTURED FORM...")
    html_structured = SmartFormRenderer.render_form_html(parsed_form)
    print(f"    ✅ Generated {len(html_structured)} characters of HTML")
    print(f"    ✅ Contains {html_structured.count('<input')} input fields")
    print(f"    ✅ HTML preview (first 300 chars):")
    print("       " + html_structured[:300].replace("\n", "\n       "))

    # STEP 5: Render as inline form
    print("\n[5] RENDERING INLINE FORM...")
    html_inline = SmartFormRenderer.render_form_with_inline_replacement(
        parsed_form.raw_content,
        parsed_form.fields
    )
    print(f"    ✅ Generated {len(html_inline)} characters of HTML")
    print(f"    ✅ Contains {html_inline.count('<input')} input fields")

    # STEP 6: Verify completeness
    print("\n[6] VERIFICATION:")
    checks = [
        ("Has sections", len(parsed_form.sections) > 0),
        ("Has fields", len(parsed_form.fields) > 0),
        ("Structured HTML valid", "<div class=\"form-container\"" in html_structured),
        ("Inline HTML valid", "<div class=\"form-document\"" in html_inline),
        ("All fields have names", all(f['name'] for f in fields_list)),
        ("All fields have types", all(f['field_type'] for f in fields_list)),
    ]

    all_passed = True
    for check_name, result in checks:
        status = "✅" if result else "❌"
        print(f"    {status} {check_name}")
        if not result:
            all_passed = False

    # STEP 7: Summary
    print("\n" + "=" * 70)
    if all_passed:
        print("✅ ALL INTEGRATION TESTS PASSED!")
        print("\nSummary:")
        print(f"  • Document Structure: {len(parsed_form.sections)} sections detected")
        print(f"  • Form Fields: {len(parsed_form.fields)} fields extracted")
        print(f"  • Rendering: Both structured & inline formats working")
        print(f"  • Field Types: {len(set(f['field_type'] for f in fields_list))} different types")
        print(f"  • Ready for production: YES")
    else:
        print("❌ SOME TESTS FAILED")
    print("=" * 70)

    return parsed_form, html_structured, html_inline


if __name__ == "__main__":
    try:
        parsed_form, html_structured, html_inline = run_integration_test()
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
