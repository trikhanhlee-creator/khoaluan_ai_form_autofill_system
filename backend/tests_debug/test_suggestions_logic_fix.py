"""Test suggestions logic fix"""
import requests
import json
import time

BASE_URL = "http://127.0.0.1:8000"

def test_suggestions_fix():
    """Test suggestions logic fix"""
    print("=" * 70)
    print("TEST: SUGGESTIONS LOGIC FIX")
    print("=" * 70)
    
    USER_ID = 1
    FORM_ID = 1
    
    # Step 1: Upload Word file
    print("\n[Step 1] Upload Word file with fields")
    print("-" * 70)
    
    test_file = "test_suggestions_fix.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("Họ và tên.....")
    doc.add_paragraph("Lớp")
    doc.add_paragraph("Trường(...)")
    doc.add_paragraph("Địa chỉ\"\"\"")
    doc.save(test_file)
    
    with open(test_file, "rb") as f:
        files = {"file": (test_file, f)}
        params = {"user_id": USER_ID}
        response = requests.post(f"{BASE_URL}/api/word/upload", files=files, params=params)
    
    if response.status_code != 200:
        print(f"❌ Upload failed: {response.json()}")
        return False
    
    result = response.json()
    template_id = result["template_id"]
    print(f"✅ Uploaded template: ID={template_id}")
    print(f"   Fields: {result['fields_count']}")
    for field in result["fields"]:
        print(f"   - {field['label']} ({field['name']})")
    
    # Step 2: Get template detail to verify field_ids are created
    print("\n[Step 2] Get template detail to verify field_ids")
    print("-" * 70)
    
    response = requests.get(
        f"{BASE_URL}/api/word/template/{template_id}",
        params={"user_id": USER_ID, "form_id": FORM_ID}
    )
    
    if response.status_code != 200:
        print(f"❌ Get template failed: {response.json()}")
        return False
    
    template_data = response.json()
    print(f"✅ Got template: {template_data['name']}")
    print(f"   Form ID: {template_data['form_id']}")
    print(f"   Fields with IDs:")
    
    field_ids = {}
    for idx, field in enumerate(template_data["fields"]):
        field_id = field.get("field_id")
        field_name = field.get("name")
        field_label = field.get("label")
        print(f"   [{idx}] {field_label:<20} → name={field_name:<20} field_id={field_id}")
        field_ids[field_name] = field_id
    
    # Verify all fields have valid IDs
    invalid_fields = [name for name, fid in field_ids.items() if not fid or fid < 0]
    if invalid_fields:
        print(f"\n❌ ERROR: Fields with invalid IDs: {invalid_fields}")
        return False
    
    print(f"\n✅ All fields have valid IDs")
    
    # Step 3: Get suggestions for each field (should be empty - no values yet)
    print("\n[Step 3] Get suggestions (should be empty initially)")
    print("-" * 70)
    
    for field_name, field_id in field_ids.items():
        response = requests.get(
            f"{BASE_URL}/api/suggestions",
            params={"user_id": USER_ID, "field_id": field_id}
        )
        data = response.json()
        print(f"   {field_name:<20} (id={field_id}): {len(data.get('suggestions', []))} suggestions")
    
    # Step 4: Save some entries for each field
    print("\n[Step 4] Save entries for each field")
    print("-" * 70)
    
    test_values = {
        "họ_và_tên": ["Nguyễn Văn A", "Trần Thị B"],
        "lớp": ["10A1", "10A2"],
        "trường": ["THPT ABC", "THPT XYZ"],
        "địa_chỉ": ["Hà Nội", "TP.HCM"]
    }
    
    for field_name, field_id in field_ids.items():
        if field_name not in test_values:
            continue
        
        for value in test_values[field_name]:
            response = requests.post(
                f"{BASE_URL}/api/suggestions/save",
                params={
                    "user_id": USER_ID,
                    "field_id": field_id,
                    "form_id": FORM_ID,
                    "value": value
                }
            )
            if response.status_code == 200:
                print(f"   ✅ Saved '{value}' for {field_name} (field_id={field_id})")
            else:
                print(f"   ❌ Failed to save '{value}' for {field_name}: {response.status_code}")
    
    # Step 5: Get suggestions for each field (should show saved values)
    print("\n[Step 5] Get suggestions (should show saved values ONLY FOR THAT FIELD)")
    print("-" * 70)
    
    all_pass = True
    for field_name, field_id in field_ids.items():
        response = requests.get(
            f"{BASE_URL}/api/suggestions",
            params={"user_id": USER_ID, "field_id": field_id}
        )
        data = response.json()
        suggestions = data.get("suggestions", [])
        
        print(f"\n   {field_name} (field_id={field_id}):")
        print(f"   Entry count: {data.get('entry_count')}")
        print(f"   Suggestions: {len(suggestions)}")
        
        for sug in suggestions:
            print(f"      - {sug.get('value')} (freq: {sug.get('frequency')})")
        
        # Verify suggestions match expected values for this field
        expected = test_values.get(field_name, [])
        suggested_values = [s.get("value") for s in suggestions]
        
        for exp_val in expected:
            if exp_val in suggested_values:
                print(f"      ✅ Expected value '{exp_val}' found")
            else:
                print(f"      ❌ Expected value '{exp_val}' NOT found")
                all_pass = False
        
        # Verify NO suggestions from other fields
        for other_field_name, other_values in test_values.items():
            if other_field_name == field_name:
                continue
            for other_val in other_values:
                if other_val in suggested_values:
                    print(f"      ❌ CROSS-FIELD ERROR: Found '{other_val}' from {other_field_name}!")
                    all_pass = False
    
    # Cleanup
    import os
    if os.path.exists(test_file):
        os.remove(test_file)
    
    print("\n" + "=" * 70)
    if all_pass:
        print("✅ TEST PASSED - Suggestions are correctly filtered by field!")
    else:
        print("❌ TEST FAILED - Suggestions logic still has issues!")
    print("=" * 70)
    
    return all_pass

if __name__ == "__main__":
    print("Starting suggestions logic test...\n")
    print("Make sure server is running: python run.py\n")
    time.sleep(1)
    
    try:
        test_suggestions_fix()
    except requests.exceptions.ConnectionError:
        print("❌ Cannot connect to server!")
        print("Make sure backend is running on http://127.0.0.1:8000")
