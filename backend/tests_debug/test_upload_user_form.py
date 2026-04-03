"""Test upload file giống file từ user"""
import requests
import json
from docx import Document
import os
import time

BASE_URL = "http://127.0.0.1:8000"

def test_upload_user_form():
    """Test upload file giống file từ user (với các field có separators khác nhau)"""
    test_file = "test_user_form.docx"
    
    # Tạo file Word giống file user
    doc = Document()
    
    # Thêm các trường như trong ảnh user
    doc.add_paragraph("Họ và tên.....")
    doc.add_paragraph("Lớp")
    doc.add_paragraph("Trường(...)")
    doc.add_paragraph("Địa chỉ\"\"\"")
    
    doc.save(test_file)
    print(f"✓ Created test file: {test_file}\n")
    
    try:
        # Upload file
        with open(test_file, "rb") as f:
            files = {"file": (test_file, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            params = {"user_id": 5}
            
            print(f"📤 Uploading file from user...")
            response = requests.post(
                f"{BASE_URL}/api/word/upload",
                files=files,
                params=params
            )
        
        print(f"Status Code: {response.status_code}")
        result = response.json()
        
        if response.status_code == 200:
            print("✅ UPLOAD SUCCESSFUL!\n")
            print(f"Template ID: {result.get('template_id')}")
            print(f"Fields Count: {result.get('fields_count')}\n")
            
            print("📋 Fields Extracted:")
            print("-" * 60)
            for idx, field in enumerate(result.get('fields', []), 1):
                print(f"{idx}. {field.get('label'):<30} → {field.get('name'):<20} ({field.get('field_type')})")
            print("-" * 60)
            
            print(f"\n✅ Form ready to use!")
            return True
        else:
            print(f"\n❌ Upload failed!")
            print(f"Error: {result.get('detail')}")
            return False
    
    finally:
        # Cleanup
        if os.path.exists(test_file):
            os.remove(test_file)
            print(f"\n✓ Cleaned up test file")

if __name__ == "__main__":
    print("=" * 70)
    print("TEST: UPLOAD FILE FROM USER")
    print("=" * 70)
    print()
    
    print("Waiting for server startup...\n")
    time.sleep(1)
    
    try:
        test_upload_user_form()
    except requests.exceptions.ConnectionError:
        print("❌ Cannot connect to server at http://127.0.0.1:8000")
        print("   Make sure backend server is running!")
