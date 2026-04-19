#!/usr/bin/env python3
"""Regression checks for parsing preformatted Word forms."""

import os
import sys
import tempfile

sys.path.insert(0, '.')

from docx import Document

from app.services.file_parser import FileParserFactory


def _make_temp_docx(build_fn):
    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    try:
        doc = Document()
        build_fn(doc)
        doc.save(path)
        return path
    except Exception:
        if os.path.exists(path):
            os.remove(path)
        raise


def _parse_labels(path: str) -> list[str]:
    parser = FileParserFactory.create_parser(path)
    fields = parser.parse()
    return [field.label for field in fields]


def test_preformatted_paragraph_lines() -> None:
    def build(doc: Document) -> None:
        doc.add_paragraph('ĐƠN XIN VIỆC')
        doc.add_paragraph('Kính gửi: Ban lãnh đạo và phòng nhân sự Công ty ..............................')
        doc.add_paragraph('Tôi tên là: ....................................................................')
        doc.add_paragraph('Sinh ngày: .....................................................................')
        doc.add_paragraph('Chỗ ở hiện nay: ................................................................')
        doc.add_paragraph('Thông qua website, tôi biết công ty có nhu cầu tuyển dụng vị trí...............')
        doc.add_paragraph('Tôi đã tốt nghiệp loại ....... tại trường ......................................')

    path = _make_temp_docx(build)
    try:
        labels = _parse_labels(path)
    finally:
        if os.path.exists(path):
            os.remove(path)

    expected = {
        'Công ty ứng tuyển',
        'Họ và tên',
        'Ngày sinh',
        'Địa chỉ hiện tại',
        'Vị trí ứng tuyển',
        'Xếp loại tốt nghiệp',
        'Trường tốt nghiệp',
    }

    missing = expected.difference(labels)
    assert not missing, f'Missing labels: {sorted(missing)} | got={labels}'


def test_preformatted_table_lines() -> None:
    def build(doc: Document) -> None:
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = 'Email'
        table.cell(0, 1).text = '......................'
        table.cell(1, 0).text = 'Ngày sinh: ...............'
        table.cell(1, 1).text = ''
        table.cell(2, 0).text = 'Giới tính'
        table.cell(2, 1).text = '[ ] Nam   [ ] Nữ'

    path = _make_temp_docx(build)
    try:
        labels = _parse_labels(path)
    finally:
        if os.path.exists(path):
            os.remove(path)

    expected = {'Email', 'Ngày sinh', 'Giới tính'}
    missing = expected.difference(labels)
    assert not missing, f'Missing labels: {sorted(missing)} | got={labels}'


def test_skip_signature_footer_date_line() -> None:
    def build(doc: Document) -> None:
        doc.add_paragraph('Tôi tên là: ....................................................................')
        doc.add_paragraph('......., ngày.... tháng.... năm ....')
        doc.add_paragraph('Người viết đơn')

    path = _make_temp_docx(build)
    try:
        labels = _parse_labels(path)
    finally:
        if os.path.exists(path):
            os.remove(path)

    assert 'Ngày nộp đơn' not in labels, f'Signature date footer should not become a field | got={labels}'


if __name__ == '__main__':
    test_preformatted_paragraph_lines()
    test_preformatted_table_lines()
    test_skip_signature_footer_date_line()
    print('All formatted Word parser regression checks passed.')
