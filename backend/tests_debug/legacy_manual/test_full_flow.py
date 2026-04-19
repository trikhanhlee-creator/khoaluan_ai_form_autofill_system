#!/usr/bin/env python3
"""Full flow test: upload, retrieve, check labels"""
import sys
import json
sys.path.insert(0, '.')
import requests
import time
from docx import Document

# Create a test file with problematic characters
doc = Document()
doc.add_paragraph('Họ và tên.....')
doc.add_paragraph('Lớp')  
doc.add_paragraph('Trường(...)')
doc.add_paragraph('Địa chỉ"""')  # 3 quotes

test_file = 'test_debug_form.docx'
doc.save(test_file)
print(f"Created test file with 3 quotes at end of last field\n")

# Wait for server
time.sleep(0.5)

# Upload
print("1. Uploading file...")
with open(test_file, 'rb') as f:
    response = requests.post(
        'http://localhost:8000/api/word/upload?user_id=1',
        files={'file': f}
    )
result = response.json()
template_id = result['template_id']
print(f"   Template ID: {template_id}\n")

# Check raw upload response
print("2. Checking upload response fields:")
for field in result['fields']:
    print(f"   Label: {field['label']}")

# Retrieve template
print("\n3. Retrieving template...")
response = requests.get(
    f'http://localhost:8000/api/word/template/{template_id}?user_id=1&form_id=1'
)
template = response.json()

print("4. Checking retrieved template fields:")
for field in template['fields']:
    label = field.get('label', '')
    has_quotes = '"' in label
    print(f"   Label: {label:20} | Has quotes: {has_quotes}")

# Save to file for inspection
output = {
    'upload_response_fields': result['fields'],
    'template_response_fields': template['fields']
}

with open('flow_test_output.json', 'w', encoding='utf-8') as f:
    json.dump(output, f, indent=2, ensure_ascii=False)

print("\nResults saved to flow_test_output.json")

# Cleanup
import os
if os.path.exists(test_file):
    os.remove(test_file)
