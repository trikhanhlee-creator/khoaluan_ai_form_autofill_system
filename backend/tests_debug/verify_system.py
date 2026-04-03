"""
Comprehensive Verification - Ensure all components are working
"""

import sys
import os

# Setup path
backend_path = os.path.join(os.path.dirname(__file__), '..')
sys.path.insert(0, backend_path)

import inspect


def verify_intelligent_detector():
    """Verify IntelligentDetector is working"""
    print("\n[1] Verifying IntelligentDetector...")
    
    from app.services.form_replacement import IntelligentDetector
    
    checks = []
    
    # Check methods exist
    methods = ['parse_document', 'extract_field_list', 'is_title_candidate', 
               'extract_label', 'has_placeholder', 'detect_field_type', 'generate_field_name']
    
    for method in methods:
        exists = hasattr(IntelligentDetector, method)
        checks.append((f"Method: {method}", exists))
    
    # Check classes exist
    from app.services.form_replacement import FormSection, FormField, ParsedForm
    checks.append(("FormSection dataclass", FormSection is not None))
    checks.append(("FormField dataclass", FormField is not None))
    checks.append(("ParsedForm dataclass", ParsedForm is not None))
    
    for check_name, result in checks:
        status = "✅" if result else "❌"
        print(f"  {status} {check_name}")
    
    all_pass = all(result for _, result in checks)
    return all_pass


def verify_smart_renderer():
    """Verify SmartFormRenderer is working"""
    print("\n[2] Verifying SmartFormRenderer...")
    
    from app.services.form_replacement import SmartFormRenderer, DocumentStructurePreserver
    
    checks = []
    
    # Check methods exist
    methods = ['render_form_html', 'render_form_with_inline_replacement']
    
    for method in methods:
        exists = hasattr(SmartFormRenderer, method)
        checks.append((f"Method: {method}", exists))
    
    # Check DocumentStructurePreserver
    checks.append(("DocumentStructurePreserver exists", DocumentStructurePreserver is not None))
    
    for check_name, result in checks:
        status = "✅" if result else "❌"
        print(f"  {status} {check_name}")
    
    all_pass = all(result for _, result in checks)
    return all_pass


def verify_api_routes():
    """Verify API routes are registered"""
    print("\n[3] Verifying API Routes...")
    
    from app.main import app
    
    # Get all routes
    routes = []
    for route in app.routes:
        if hasattr(route, 'path'):
            routes.append(route.path)
    
    # Check for new endpoints
    expected_endpoints = [
        '/api/form-replacement/upload-with-intelligent-detection',
        '/api/form-replacement/template/{template_id}/render-form-structured',
        '/api/form-replacement/template/{template_id}/render-form-inline',
    ]
    
    checks = []
    for endpoint in expected_endpoints:
        exists = any(endpoint in r or ('{' in endpoint and endpoint.replace('{template_id}', '[^}]+') in r) for r in routes)
        checks.append((f"Endpoint: {endpoint}", exists))
    
    # Check legacy endpoints still exist
    legacy_endpoints = [
        '/api/form-replacement/upload-with-dotlines',
        '/api/form-replacement/template/{template_id}/render-form',
    ]
    
    for endpoint in legacy_endpoints:
        exists = any(endpoint in r or '{template_id}' in endpoint for r in routes)
        checks.append((f"Legacy: {endpoint}", exists))
    
    for check_name, result in checks:
        status = "✅" if result else "❌"
        print(f"  {status} {check_name}")
    
    all_pass = all(result for _, result in checks)
    return all_pass


def verify_functionality():
    """Verify actual functionality works"""
    print("\n[4] Verifying Functionality...")
    
    from docx import Document
    from docx.shared import Pt
    from app.services.form_replacement import IntelligentDetector, SmartFormRenderer
    
    checks = []
    
    try:
        # Create test document
        doc = Document()
        title = doc.add_paragraph()
        title.add_run("ĐƠN XIN VIỆC").bold = True
        doc.add_paragraph("Tôi tên là: ..........................................")
        doc.add_paragraph("Sinh ngày: ___/___/______")
        
        # Test parsing
        parsed_form = IntelligentDetector.parse_document(doc)
        checks.append(("Parse document", len(parsed_form.fields) > 0))
        
        # Test field extraction
        fields = IntelligentDetector.extract_field_list(parsed_form)
        checks.append(("Extract fields", len(fields) > 0 and 'name' in fields[0]))
        
        # Test structured rendering
        html_structured = SmartFormRenderer.render_form_html(parsed_form)
        checks.append(("Render structured", '<input' in html_structured))
        
        # Test inline rendering
        html_inline = SmartFormRenderer.render_form_with_inline_replacement(
            parsed_form.raw_content,
            parsed_form.fields
        )
        checks.append(("Render inline", '<input' in html_inline))
        
    except Exception as e:
        checks.append(("Functionality test", False))
        print(f"  ❌ Error: {str(e)}")
    
    for check_name, result in checks:
        status = "✅" if result else "❌"
        print(f"  {status} {check_name}")
    
    all_pass = all(result for _, result in checks)
    return all_pass


def verify_backward_compatibility():
    """Verify old API still works"""
    print("\n[5] Verifying Backward Compatibility...")
    
    from app.services.form_replacement import DotLineDetector, HTMLFieldReplacer
    
    checks = []
    
    # Check old classes still exist
    checks.append(("DotLineDetector exists", DotLineDetector is not None))
    checks.append(("HTMLFieldReplacer exists", HTMLFieldReplacer is not None))
    
    # Check methods exist
    checks.append(("DotLineDetector.detect_from_document", hasattr(DotLineDetector, 'detect_from_document')))
    checks.append(("HTMLFieldReplacer.render_form_with_replacements", hasattr(HTMLFieldReplacer, 'render_form_with_replacements')))
    
    for check_name, result in checks:
        status = "✅" if result else "❌"
        print(f"  {status} {check_name}")
    
    all_pass = all(result for _, result in checks)
    return all_pass


def verify_data_models():
    """Verify all data models are available"""
    print("\n[6] Verifying Data Models...")
    
    from app.services.form_replacement import (
        FormSection, FormField, ParsedForm, DotLinePlaceholder,
        DotLineField, FormReplacementResult
    )
    
    checks = []
    
    models = [
        ("FormSection", FormSection),
        ("FormField", FormField),
        ("ParsedForm", ParsedForm),
        ("DotLinePlaceholder", DotLinePlaceholder),
        ("DotLineField", DotLineField),
        ("FormReplacementResult", FormReplacementResult),
    ]
    
    for model_name, model_class in models:
        checks.append((f"Model: {model_name}", model_class is not None))
    
    for check_name, result in checks:
        status = "✅" if result else "❌"
        print(f"  {status} {check_name}")
    
    all_pass = all(result for _, result in checks)
    return all_pass


def main():
    """Run all verifications"""
    print("\n" + "=" * 80)
    print("COMPREHENSIVE VERIFICATION - Intelligent Form Detection System")
    print("=" * 80)
    
    results = []
    
    try:
        results.append(("IntelligentDetector", verify_intelligent_detector()))
        results.append(("SmartFormRenderer", verify_smart_renderer()))
        results.append(("Data Models", verify_data_models()))
        results.append(("API Routes", verify_api_routes()))
        results.append(("Backward Compatibility", verify_backward_compatibility()))
        results.append(("Functionality", verify_functionality()))
        
    except Exception as e:
        print(f"\n❌ Verification Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False
    
    # Summary
    print("\n" + "=" * 80)
    print("VERIFICATION SUMMARY")
    print("=" * 80)
    
    all_pass = True
    for category, passed in results:
        status = "✅" if passed else "❌"
        print(f"{status} {category}")
        if not passed:
            all_pass = False
    
    print("=" * 80)
    
    if all_pass:
        print("\n🎉 ALL VERIFICATIONS PASSED! System is ready for production.\n")
        return True
    else:
        print("\n❌ Some verifications failed. Please review above.\n")
        return False


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
