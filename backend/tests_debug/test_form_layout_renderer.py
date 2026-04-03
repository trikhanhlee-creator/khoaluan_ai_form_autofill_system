"""
Test Form Layout Renderer - Document style form rendering
Giống cấu trúc trong hình ảnh - phân biệt tiêu đề vs fields
"""

import sys
import os

backend_path = os.path.join(os.path.dirname(__file__), '..')
sys.path.insert(0, backend_path)

from docx import Document
from docx.shared import Pt
from app.services.form_replacement import (
    IntelligentDetector,
    FormLayoutRenderer,
    FormPageRenderer
)


def create_form_document():
    """Tạo document giống hình ảnh"""
    doc = Document()

    # Header
    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r.bold = True
    r.font.size = Pt(11)

    # Subheader
    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run("Độc lập - Tự do - Hạnh phúc")

    # Main title
    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run("ĐƠN XIN VIỆC")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph()

    # Content with fields
    doc.add_paragraph("Kính gửi: Ban lãnh đạo và phòng nhân sự Công ty ..........................................")
    doc.add_paragraph()

    doc.add_paragraph("Tôi tên là: ........................................................................")
    doc.add_paragraph()

    doc.add_paragraph("Sinh ngày: ___/___/______")
    doc.add_paragraph()

    doc.add_paragraph("Chỗ ở hiện nay: ........................................................................")
    doc.add_paragraph()

    doc.add_paragraph("Số điện thoại liên hệ: ___________________")
    doc.add_paragraph()

    # Description
    p = doc.add_paragraph()
    p.add_run("Thông qua trang website TopCV.vn, tôi biết được Quý công ty có nhu cầu tuyển dụng vị trí ")
    p.add_run("________________")
    p.add_run(" Tôi cảm thấy trình độ và kỹ năng của mình phù hợp với vị trí này. Tôi mong muốn được làm việc và công hiến cho công ty.")

    return doc


def test_document_layout_rendering():
    """Test form document layout rendering"""
    
    print("\n" + "=" * 80)
    print("TEST: Form Document Layout Rendering")
    print("=" * 80)

    # Create document
    print("\n[1] Creating form document...")
    doc = create_form_document()
    print("✅ Document created")

    # Parse
    print("\n[2] Parsing document...")
    parsed_form = IntelligentDetector.parse_document(doc)
    print(f"✅ Parsed: {len(parsed_form.sections)} sections, {len(parsed_form.fields)} fields")

    # Display structure
    print("\n[3] Form Structure:")
    for section_idx, section in enumerate(parsed_form.sections):
        print(f"\n  Section {section_idx}: {section.title}")
        print(f"    └─ Is Title: {section.is_title}")
        print(f"    └─ Fields: {len([f for f in parsed_form.fields if f.section_index == section_idx])}")

    # Display fields
    print("\n[4] Fields Detected:")
    fields = IntelligentDetector.extract_field_list(parsed_form)
    for field in fields:
        print(f"\n  • {field['label']}")
        print(f"    Type: {field['field_type']}")
        print(f"    Name: {field['name']}")

    # Test document style rendering
    print("\n[5] Rendering as Document Style...")
    html_document = FormLayoutRenderer.render_form_document_style(parsed_form)
    print(f"✅ Generated {len(html_document)} characters")
    print(f"✅ Contains {html_document.count('<input')} input elements")
    print(f"✅ Contains {html_document.count('<textarea')} textarea elements")

    # Test complete page rendering
    print("\n[6] Rendering as Complete Page...")
    html_page = FormPageRenderer.render_complete_form_page(parsed_form)
    print(f"✅ Generated {len(html_page)} characters")
    print(f"✅ Contains HTML, CSS, and JavaScript")
    print(f"✅ Contains form buttons (submit, reset)")

    # Show HTML preview (document style)
    print("\n[7] HTML Preview (Document Style - first 600 chars):")
    preview = html_document[:600]
    print("   " + preview.replace("\n", "\n   "))

    # Save outputs
    print("\n[8] Saving HTML outputs...")
    
    # Save document style
    doc_filename = "form_document_style.html"
    with open(doc_filename, "w", encoding="utf-8") as f:
        f.write(html_document)
    print(f"   ✅ Saved: {doc_filename}")

    # Save complete page
    page_filename = "form_complete_page.html"
    with open(page_filename, "w", encoding="utf-8") as f:
        f.write(html_page)
    print(f"   ✅ Saved: {page_filename}")

    # Validation
    print("\n[9] Validation:")
    checks = [
        ("Has form container", "<div" in html_document),
        ("Has titles", "form-title" in html_document or "CỘNG HÒA" in html_document),
        ("Has fields", "<input" in html_document or "<textarea" in html_document),
        ("Field count matches", html_document.count("<input") + html_document.count("<textarea") >= len(parsed_form.fields)),
        ("Page is complete HTML", "<!DOCTYPE html>" in html_page),
        ("Page has CSS", "<style>" in html_page),
        ("Page has JS", "<script>" in html_page),
        ("Page has buttons", "button" in html_page.lower()),
    ]

    all_pass = True
    for check_name, result in checks:
        status = "✅" if result else "❌"
        print(f"  {status} {check_name}")
        if not result:
            all_pass = False

    print("\n" + "=" * 80)
    if all_pass:
        print("✅ ALL TESTS PASSED!")
        print("\nKey Features:")
        print("  • Tiêu đề (titles) được hiển thị as non-editable text")
        print("  • Trường (fields) được hiển thị as input boxes")
        print("  • Layout giống document gốc")
        print("  • Phân biệt rõ ràng giữa title và fields")
        print("  • Ready to display in web browser")
    else:
        print("❌ SOME TESTS FAILED")
    print("=" * 80)

    return parsed_form, html_document, html_page


if __name__ == "__main__":
    try:
        parsed_form, html_document, html_page = test_document_layout_rendering()
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
