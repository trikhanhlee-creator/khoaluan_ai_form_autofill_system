"""Test upload file Word với fallback logic"""
import requests
import json
from docx import Document
import os
import time

# Chờ server startup
time.sleep(2)

BASE_URL = "http://127.0.0.1:8000"
UPLOAD_DIR = "uploads"

# Tạo file Word test cơ bản (không có field cấu trúc)
def create_test_word_file(filename):
    """Tạo file Word đơn giản không có fields cấu trúc"""
    doc = Document()
    doc.add_paragraph("Đây là một tài liệu test")
    doc.add_paragraph("Nó chứa một số thông tin chung")
    doc.add_paragraph("Nhưng không có fields định dạng chuẩn")
    doc.add_paragraph("Hãy kiểm tra xem có thể tạo form được không")
    
    doc.save(filename)
    print(f"✓ Created test file: {filename}")

# Test upload avec fallback
def test_upload_with_fallback():
    """Test upload file không có fields cấu trúc"""
    test_file = "test_simple_form.docx"
    
    # Tạo test file
    create_test_word_file(test_file)
    
    try:
        # Upload file
        with open(test_file, "rb") as f:
            files = {"file": (test_file, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            params = {"user_id": 1}
            
            print(f"\n📤 Uploading {test_file}...")
            response = requests.post(
                f"{BASE_URL}/api/word/upload",
                files=files,
                params=params
            )
        
        print(f"Status Code: {response.status_code}")
        result = response.json()
        print(f"Response: {json.dumps(result, indent=2, ensure_ascii=False)}")
        
        if response.status_code == 200:
            print("\n✅ SUCCESS! Upload thành công ngay cả khi không có fields cấu trúc")
            print(f"   - Template ID: {result.get('template_id')}")
            print(f"   - Fields tạo được: {result.get('fields_count')}")
            print(f"   - Auto generated: {result.get('auto_generated_fields')}")
            return True
        else:
            print(f"\n❌ FAILED! Error: {result.get('detail')}")
            return False
    
    finally:
        # Cleanup
        if os.path.exists(test_file):
            os.remove(test_file)
            print(f"✓ Cleaned up {test_file}")

# Test upload file với structured fields
def test_upload_with_structure():
    """Test upload file có fields định dạng chuẩn"""
    test_file = "test_structured_form.docx"
    
    # Tạo test file với fields chuẩn
    doc = Document()
    doc.add_paragraph("FORM ĐĂNG KÝ")
    doc.add_paragraph("Họ và tên:")
    doc.add_paragraph("Email:")
    doc.add_paragraph("Số điện thoại:")
    doc.add_paragraph("Sinh năm:")
    
    doc.save(test_file)
    print(f"✓ Created structured test file: {test_file}")
    
    try:
        # Upload file
        with open(test_file, "rb") as f:
            files = {"file": (test_file, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            params = {"user_id": 2}
            
            print(f"\n📤 Uploading {test_file}...")
            response = requests.post(
                f"{BASE_URL}/api/word/upload",
                files=files,
                params=params
            )
        
        print(f"Status Code: {response.status_code}")
        result = response.json()
        print(f"Response: {json.dumps(result, indent=2, ensure_ascii=False)}")
        
        if response.status_code == 200:
            print("\n✅ SUCCESS! Upload thành công")
            print(f"   - Template ID: {result.get('template_id')}")
            print(f"   - Fields tạo được: {result.get('fields_count')}")
            for field in result.get('fields', []):
                print(f"     • {field.get('name')} ({field.get('field_type')}): {field.get('label')}")
            return True
        else:
            print(f"\n❌ FAILED! Error: {result.get('detail')}")
            return False
    
    finally:
        # Cleanup
        if os.path.exists(test_file):
            os.remove(test_file)
            print(f"✓ Cleaned up {test_file}")

if __name__ == "__main__":
    print("=" * 60)
    print("TEST UPLOAD FILE WORD VỚI FALLBACK LOGIC")
    print("=" * 60)
    
    print("\n[TEST 1] Upload file không có fields cấu trúc (fallback)")
    success1 = test_upload_with_fallback()
    
    print("\n" + "=" * 60)
    print("\n[TEST 2] Upload file có structured fields")
    success2 = test_upload_with_structure()
    
    print("\n" + "=" * 60)
    if success1 and success2:
        print("✅ TẤT CẢ TESTS PASS!")
    else:
        print("❌ CÓ TESTS FAILED")
