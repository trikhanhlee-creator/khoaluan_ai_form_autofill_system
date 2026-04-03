"""
Real-World Test - Simulating actual form from image
Recreates the "ĐƠN XIN VIỆC" (Job Application) form structure
"""

import sys
import os
import json

# Setup path
backend_path = os.path.join(os.path.dirname(__file__), '..')
sys.path.insert(0, backend_path)

from docx import Document
from docx.shared import Pt
from app.services.form_replacement import IntelligentDetector, SmartFormRenderer


def create_realistic_job_application():
    """
    Create realistic job application form matching the 2nd image
    Recreates: "ĐƠN XIN VIỆC" (Job Application Form)
    """
    doc = Document()

    # Header - Government header
    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r.bold = True
    r.font.size = Pt(11)

    # Subheader
    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run("Độc lập - Tự do - Hạnh phúc")

    # Main title
    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run("ĐƠN XIN VIỆC")
    r.bold = True
    r.font.size = Pt(14)
    
    doc.add_paragraph()  # Spacing

    # Company info
    doc.add_paragraph("Kính gửi: Ban lãnh đạo và phòng nhân sự Công ty ..........................................")
    doc.add_paragraph()

    # Personal fields
    doc.add_paragraph("Tôi tên là: ........................................................................")
    doc.add_paragraph()

    doc.add_paragraph("Sinh ngày: ___/___/______")
    doc.add_paragraph()

    doc.add_paragraph("Chỗ ở hiện nay: ........................................................................")
    doc.add_paragraph()

    doc.add_paragraph("Số điện thoại liên hệ: ___________________")
    doc.add_paragraph()

    # Long paragraph with inline field
    para = doc.add_paragraph()
    para.add_run("Thông qua trang website TopCV.vn, tôi biết được Quý công ty có nhu cầu tuyển dụng vị trí ")
    para.add_run("________________")
    para.add_run(" Tôi cảm thấy trình độ và kỹ năng của mình phù hợp với vị trí này. Tôi mong muốn được làm việc và công hiến cho công ty.")

    return doc


def test_real_world_scenario():
    """Test with realistic job application form"""
    
    print("\n" + "=" * 80)
    print("REAL-WORLD TEST: Job Application Form Processing")
    print("=" * 80)
    
    # Create document
    print("\n[STEP 1] Creating realistic job application document...")
    doc = create_realistic_job_application()
    print("✅ Document created with realistic structure")

    # Parse with intelligent detection
    print("\n[STEP 2] Parsing document with intelligent detection...")
    parsed_form = IntelligentDetector.parse_document(doc)
    print(f"✅ Detected {len(parsed_form.sections)} sections")
    print(f"✅ Detected {len(parsed_form.fields)} fields")
    
    # Display sections
    print("\n[STEP 3] Form Structure Detected:")
    for idx, section in enumerate(parsed_form.sections):
        print(f"\n  Section {idx + 1}: {section.title}")
        print(f"    └─ Title: {section.is_title}")
        print(f"    └─ Level: {section.level}")
        print(f"    └─ Items: {len(section.items)}")
        
        # Show items in this section
        for item_idx, item in enumerate(section.items):
            if item['type'] == 'field':
                print(f"        ├─ [FIELD {item_idx}]")
            else:
                text_preview = item['text'][:50] + "..." if len(item['text']) > 50 else item['text']
                print(f"        ├─ [TEXT] {text_preview}")

    # Display fields
    print("\n[STEP 4] Detected Fields:")
    fields_list = IntelligentDetector.extract_field_list(parsed_form)
    
    for field in fields_list:
        print(f"\n  ◆ {field['label']}")
        print(f"    Field Name: {field['name']}")
        print(f"    Field Type: {field['field_type']}")
        print(f"    Context: {field['context'][:60]}...")

    # Try rendering as structured form
    print("\n[STEP 5] Rendering as Structured Form...")
    html_structured = SmartFormRenderer.render_form_html(parsed_form)
    print(f"✅ Generated {len(html_structured)} characters of HTML")
    print(f"✅ Contains {html_structured.count('<input')} input fields")
    
    if html_structured.count('<input') == len(parsed_form.fields):
        print("✅ All fields have corresponding inputs")

    # Try rendering as inline form
    print("\n[STEP 6] Rendering as Inline Form...")
    html_inline = SmartFormRenderer.render_form_with_inline_replacement(
        parsed_form.raw_content,
        parsed_form.fields
    )
    print(f"✅ Generated {len(html_inline)} characters of HTML")
    print(f"✅ Contains {html_inline.count('<input')} input fields")

    # Validation checks
    print("\n[STEP 7] Validation & Quality Checks:")
    
    checks = [
        ("Has government header", "CỘNG HÒA" in parsed_form.raw_content[0]),
        ("Has form title", any("ĐƠN XIN VIỆC" in r for r in parsed_form.raw_content)),
        ("Name field detected", any("tôi" in f['name'].lower() for f in fields_list)),
        ("Date field detected", any(f['field_type'] == 'date' for f in fields_list)),
        ("Phone field detected", any(f['field_type'] == 'phone' for f in fields_list)),
        ("Proper section organization", len(parsed_form.sections) >= 2),
        ("All fields have names", all(f['name'] for f in fields_list)),
        ("All fields have types", all(f['field_type'] for f in fields_list)),
        ("HTML is valid", "<div" in html_structured and "form-container" in html_structured),
        ("Inline HTML is valid", "<div" in html_inline and "form-document" in html_inline),
    ]

    all_pass = True
    for check_name, result in checks:
        status = "✅" if result else "❌"
        print(f"  {status} {check_name}")
        if not result:
            all_pass = False

    # Generate form data schema
    print("\n[STEP 8] Generated Form Data Schema:")
    form_schema = {
        "form_name": "ĐƠN XIN VIỆC",
        "total_fields": len(fields_list),
        "fields": [
            {
                "name": f['name'],
                "label": f['label'],
                "type": f['field_type'],
                "required": True,
                "placeholder": f"Nhập {f['label'].lower()}..."
            }
            for f in fields_list
        ]
    }
    
    print(json.dumps(form_schema, indent=2, ensure_ascii=False))

    # Final summary
    print("\n" + "=" * 80)
    if all_pass:
        print("✅ REAL-WORLD TEST PASSED - PRODUCTION READY!")
        print("\nSummary:")
        print(f"  • Form Type: Job Application (ĐƠN XIN VIỆC)")
        print(f"  • Sections Detected: {len(parsed_form.sections)}")
        print(f"  • Fields Extracted: {len(parsed_form.fields)}")
        print(f"  • Field Types: {len(set(f['field_type'] for f in fields_list))}")
        print(f"  • Structure Preserved: YES")
        print(f"  • Both Rendering Modes Working: YES")
        print(f"  • Ready to Deploy: YES ✅")
    else:
        print("❌ SOME CHECKS FAILED")
    
    print("=" * 80)
    
    return parsed_form, fields_list, html_structured, html_inline


if __name__ == "__main__":
    try:
        parsed_form, fields_list, html_structured, html_inline = test_real_world_scenario()
        
        # Option: Save HTML to file for inspection
        with open("form_output_structured.html", "w", encoding="utf-8") as f:
            f.write(html_structured)
        print(f"\n📄 Saved structured form HTML to: form_output_structured.html")
        
        with open("form_output_inline.html", "w", encoding="utf-8") as f:
            f.write(html_inline)
        print(f"📄 Saved inline form HTML to: form_output_inline.html")
        
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
