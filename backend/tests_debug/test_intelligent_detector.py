"""
Test Intelligent Form Detector
"""

import sys
import os

# Setup path
backend_path = os.path.join(os.path.dirname(__file__), '..')
sys.path.insert(0, backend_path)

from docx import Document
from docx.shared import Pt, Inches
from app.services.form_replacement.intelligent_detector import IntelligentDetector, FormField, ParsedForm


def create_job_application_form():
    """Tạo test document giống như trong ảnh - ĐƠN XIN VIỆC"""
    doc = Document()

    # Header
    title_para = doc.add_paragraph()
    title_para.alignment = 1  # Center
    run = title_para.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    run.bold = True
    run.font.size = Pt(11)

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = 1
    run = subtitle_para.add_run("Độc lập - Tự do - Hạnh phúc")
    run.bold = True

    # Main title
    main_title = doc.add_paragraph()
    main_title.alignment = 1
    run = main_title.add_run("ĐƠN XIN VIỆC")
    run.bold = True
    run.font.size = Pt(14)

    # Content
    doc.add_paragraph("Kính gửi: Ban lãnh đạo và phòng nhân sự Công ty ..........................................")
    doc.add_paragraph()

    # Fields
    doc.add_paragraph("Tôi tên là: ........................................................................")
    doc.add_paragraph()

    doc.add_paragraph("Sinh ngày: ___/___/______  ")
    doc.add_paragraph()

    doc.add_paragraph("Chỗ ở hiện nay: ........................................................................")
    doc.add_paragraph()

    doc.add_paragraph("Số điện thoại liên hệ: ___________________")
    doc.add_paragraph()

    # Text content
    doc.add_paragraph("Thông qua trang website TopCV.vn, tôi biết được Quý công ty có nhu cầu tuyển dụng vị trí ________________ Tôi cảm thấy trình độ và kỹ năng của mình phù hợp với vị trí này. Tôi mong muốn được làm việc và công hiến cho công ty.")

    return doc


def test_intelligent_detection():
    """Test intelligent form detection"""
    print("=" * 60)
    print("TEST: Intelligent Form Detection")
    print("=" * 60)

    # Create test document
    doc = create_job_application_form()

    # Parse document
    parsed_form = IntelligentDetector.parse_document(doc)

    # Print results
    print("\n✅ SECTION DETECTION:")
    for idx, section in enumerate(parsed_form.sections):
        print(f"\n  Section {idx}:")
        print(f"    Title: {section.title[:60] if section.title else '(no title)'}")
        print(f"    Is Title: {section.is_title}")
        print(f"    Level: {section.level}")
        print(f"    Items: {len(section.items)}")
        for item in section.items:
            if item['type'] == 'field':
                print(f"      - FIELD (index {item['field_index']})")
            else:
                print(f"      - TEXT: {item['text'][:50]}")

    print("\n✅ FIELD DETECTION:")
    for idx, field in enumerate(parsed_form.fields):
        print(f"\n  Field {idx}:")
        print(f"    Label: {field.label}")
        print(f"    Name: {field.field_name}")
        print(f"    Type: {field.field_type}")
        print(f"    Section: {field.section_index}")
        print(f"    Context: {field.context[:60]}")

    print("\n" + "=" * 60)
    print(f"Total Sections: {len(parsed_form.sections)}")
    print(f"Total Fields: {len(parsed_form.fields)}")
    print(f"Total Content Lines: {len(parsed_form.raw_content)}")
    print("=" * 60)

    # Test field extraction
    fields_list = IntelligentDetector.extract_field_list(parsed_form)
    print("\n✅ EXTRACTED FIELDS:")
    for field in fields_list:
        print(f"\n  Order: {field['order']}")
        print(f"  Name: {field['name']}")
        print(f"  Label: {field['label']}")
        print(f"  Type: {field['field_type']}")

    return parsed_form


def test_smart_rendering(parsed_form: ParsedForm):
    """Test smart form rendering"""
    print("\n" + "=" * 60)
    print("TEST: Smart Form Rendering")
    print("=" * 60)

    from app.services.form_replacement.smart_form_renderer import SmartFormRenderer

    # Render form
    html = SmartFormRenderer.render_form_html(parsed_form)

    print("\n✅ RENDERED HTML (first 500 chars):")
    print(html[:500] + "..." if len(html) > 500 else html)

    print(f"\n✅ HTML Stats:")
    print(f"  Total Length: {len(html)} characters")
    print(f"  Number of <input>: {html.count('<input ')}")
    print(f"  Number of <textarea>: {html.count('<textarea')}")

    return html


if __name__ == "__main__":
    try:
        # Test detection
        parsed_form = test_intelligent_detection()

        # Test rendering
        html = test_smart_rendering(parsed_form)

        print("\n✅ ALL TESTS PASSED!")

    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
