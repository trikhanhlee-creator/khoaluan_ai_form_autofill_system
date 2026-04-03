"""Test với file Word giống file user upload"""
import requests
import json
from docx import Document
import os
import time

BASE_URL = "http://127.0.0.1:8000"

def test_real_world_form():
    """Test upload file form thực tế (giống file user)"""
    test_file = "test_real_form.docx"
    
    # Tạo file Word giống thực tế
    doc = Document()
    
    # Thêm tiêu đề
    title = doc.add_heading('HÓA ĐƠN ĐIỀU CHỈNH', 0)
    
    # Thêm nội dung không có field cấu trúc
    doc.add_paragraph('Công ty cổ phần ABC')
    doc.add_paragraph('Địa chỉ: 123 Đường ABC, Quận 1, TP.HCM')
    
    # Thêm bảng 1 dòng (không phải 2 cột field)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Row headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Mô tả'
    hdr_cells[2].text = 'Số tiền'
    
    # Data rows
    row_cells = table.rows[1].cells
    row_cells[0].text = '1'
    row_cells[1].text = 'Sản phẩm A'
    row_cells[2].text = '100.000'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = '2'
    row_cells[1].text = 'Sản phẩm B'
    row_cells[2].text = '200.000'
    
    # Thêm text sau bảng
    doc.add_paragraph('Tổng cộng: 300.000')
    doc.add_paragraph('Ghi chú: Đây là hóa đơn điều chỉnh')
    
    doc.save(test_file)
    print(f"✓ Created real-world test file: {test_file}")
    
    try:
        # Upload file
        with open(test_file, "rb") as f:
            files = {"file": (test_file, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            params = {"user_id": 3}
            
            print(f"\n📤 Uploading real-world form...")
            response = requests.post(
                f"{BASE_URL}/api/word/upload",
                files=files,
                params=params
            )
        
        print(f"\nStatus Code: {response.status_code}")
        result = response.json()
        
        if response.status_code == 200:
            print("✅ FORM UPLOADED SUCCESSFULLY!")
            print(f"\n📋 Template Details:")
            print(f"   - Template ID: {result.get('template_id')}")
            print(f"   - Template Name: {result.get('template_name')}")
            print(f"   - Fields Downloaded: {result.get('fields_count')}")
            print(f"   - Status: {result.get('message')}")
            
            if result.get('fields'):
                print(f"\n📑 Fields that can be filled:")
                for idx, field in enumerate(result.get('fields', []), 1):
                    print(f"   {idx}. {field.get('label')} (Type: {field.get('field_type')})")
            
            print(f"\n✅ Form ready to use and fill!")
            return True
        else:
            print(f"\n❌ FAILED!")
            print(f"Error: {result.get('detail')}")
            return False
    
    finally:
        # Cleanup
        if os.path.exists(test_file):
            os.remove(test_file)
            print(f"\n✓ Cleaned up test file")

if __name__ == "__main__":
    print("=" * 70)
    print("TEST: UPLOAD REAL-WORLD WORD FILE (GIỐNG FILE USER)")
    print("=" * 70)
    
    test_real_world_form()
