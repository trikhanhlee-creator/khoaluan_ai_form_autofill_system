"""Test trích xuất tên trường với loại bỏ ký tự phân tách"""
import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from docx import Document
from app.services.word_parser import WordParser
import json

def test_field_cleaning():
    """Test clean_field_label logic"""
    
    # Tạo test file Word
    test_file = "test_field_cleaning.docx"
    doc = Document()
    
    # Thêm các trường với ký tự phân tách khác nhau
    doc.add_paragraph("Họ và tên.....")
    doc.add_paragraph("Lớp")
    doc.add_paragraph("Trường(...)")
    doc.add_paragraph("Địa chỉ\"\"\"")
    doc.add_paragraph("Email:")
    doc.add_paragraph("Số điện thoại ────")
    doc.add_paragraph("Năm sinh (----)")
    doc.add_paragraph("Ghi chú [...]")
    
    doc.save(test_file)
    print(f"✓ Created test file: {test_file}\n")
    
    try:
        # Parse file
        parser = WordParser(test_file)
        fields = parser.parse()
        
        print("=" * 70)
        print("FIELD CLEANING TEST RESULTS")
        print("=" * 70)
        
        print(f"\nTotal fields extracted: {len(fields)}\n")
        
        print("Field Details:")
        print("-" * 70)
        print(f"{'#':<3} {'Label (Raw)':<25} {'Cleaned':<25} {'Name':<20}")
        print("-" * 70)
        
        for idx, field in enumerate(fields, 1):
            print(f"{idx:<3} {field.label:<25} {field.label:<25} {field.name:<20}")
        
        print("-" * 70)
        
        # Verify expected results
        print("\n✅ EXPECTED vs ACTUAL:")
        print("-" * 70)
        
        expected = {
            "Họ và tên": "họ_và_tên",
            "Lớp": "lớp",
            "Trường": "trường",
            "Địa chỉ": "địa_chỉ",
            "Email": "email",
            "Số điện thoại": "số_điện_thoại",
            "Năm sinh": "năm_sinh",
            "Ghi chú": "ghi_chú"
        }
        
        all_match = True
        for expected_label, expected_name in expected.items():
            found = False
            for field in fields:
                if field.label == expected_label and field.name == expected_name:
                    print(f"✅ '{expected_label}' → '{expected_name}' (MATCH)")
                    found = True
                    break
            
            if not found:
                # Try to find by label only
                found_label = [f for f in fields if f.label == expected_label]
                if found_label:
                    field = found_label[0]
                    print(f"⚠️  '{expected_label}' → '{field.name}' (expected: '{expected_name}')")
                    all_match = False
                else:
                    print(f"❌ '{expected_label}' NOT FOUND")
                    all_match = False
        
        print("-" * 70)
        
        if all_match:
            print("\n✅ ALL FIELDS CLEANED CORRECTLY!")
        else:
            print("\n⚠️  SOME FIELDS NEED ATTENTION")
        
        print(f"\nFull JSON Response:")
        print(json.dumps([f.to_dict() for f in fields], indent=2, ensure_ascii=False))
        
    finally:
        # Cleanup
        if os.path.exists(test_file):
            os.remove(test_file)
            print(f"\n✓ Cleaned up {test_file}")

if __name__ == "__main__":
    print("Starting field cleaning test...\n")
    test_field_cleaning()
