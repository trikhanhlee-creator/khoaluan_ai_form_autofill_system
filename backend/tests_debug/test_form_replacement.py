"""
Test Form Replacement Service
"""

import sys
import os

# Add backend directory to path
backend_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, backend_path)

from docx import Document
from app.services.form_replacement.dot_line_detector import DotLineDetector
from app.services.form_replacement.field_replacer import HTMLFieldReplacer


def create_test_doc_with_dotlines():
    """Create test document với dot-lines"""
    doc = Document()
    
    # Add title
    title = doc.add_paragraph()
    title.text = 'HÓA ĐƠN HÀNG'
    title.runs[0].bold = True
    
    # Add lines with dot-lines
    doc.add_paragraph('Họ tên khách hàng: .............................')
    doc.add_paragraph('Ngày: ___/___/______')
    doc.add_paragraph('Địa chỉ giao hàng: .....................................')
    doc.add_paragraph('Số điện thoại: ___________________')
    doc.add_paragraph('')
    doc.add_paragraph('Tôi là .......................... và tôi muốn gửi hóa đơn này')
    
    test_file = 'test_dotline_form.docx'
    doc.save(test_file)
    return test_file


def test_dot_line_detection():
    """Test detection của dot-lines"""
    print("\n" + "="*60)
    print("TEST 1: Dot-Line Detection")
    print("="*60)
    
    test_file = create_test_doc_with_dotlines()
    
    try:
        # Load document
        doc = Document(test_file)
        
        # Detect placeholders
        placeholders = DotLineDetector.detect_from_document(doc)
        print(f"✅ Found {len(placeholders)} placeholders")
        
        for idx, p in enumerate(placeholders):
            print(f"  [{idx}] {p.text[:50]}...")
            print(f"      Label: {p.label}")
            print(f"      Position: {p.dot_start_pos}-{p.dot_end_pos}")
        
        if len(placeholders) > 0:
            print("\n✅ TEST PASSED")
        else:
            print("\n⚠️ No placeholders detected")
            
    finally:
        if os.path.exists(test_file):
            os.remove(test_file)


def test_field_extraction():
    """Test extraction của fields từ placeholders"""
    print("\n" + "="*60)
    print("TEST 2: Field Extraction")
    print("="*60)
    
    test_file = create_test_doc_with_dotlines()
    
    try:
        doc = Document(test_file)
        placeholders = DotLineDetector.detect_from_document(doc)
        
        # Extract fields
        fields = DotLineDetector.extract_fields(placeholders)
        print(f"✅ Extracted {len(fields)} fields")
        
        for idx, field in enumerate(fields):
            print(f"  [{idx}] {field['label']}")
            print(f"      Name: {field['name']}")
            print(f"      Type: {field['field_type']}")
        
        if len(fields) > 0:
            print("\n✅ TEST PASSED")
        else:
            print("\n⚠️ No fields extracted")
            
    finally:
        if os.path.exists(test_file):
            os.remove(test_file)


def test_dot_line_replacement():
    """Test replacement của dot-lines"""
    print("\n" + "="*60)
    print("TEST 3: Dot-Line Replacement")
    print("="*60)
    
    test_file = create_test_doc_with_dotlines()
    
    try:
        doc = Document(test_file)
        
        # Detect & extract
        placeholders = DotLineDetector.detect_from_document(doc)
        fields = DotLineDetector.extract_fields(placeholders)
        
        print(f"✅ Detected {len(placeholders)} placeholders")
        print(f"✅ Extracted {len(fields)} fields")
        
        # Get paragraph texts
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Render form
        html_form = HTMLFieldReplacer.render_form_with_replacements(
            paragraphs,
            placeholders,
            fields
        )
        
        print(f"\n✅ Generated HTML ({len(html_form)} chars)")
        print(f"   Contains '<input': {'<input' in html_form}")
        print(f"   Contains 'dot-line-form': {'dot-line-form' in html_form}")
        print(f"   Number of inputs: {html_form.count('<input')}")
        
        if '<input' in html_form and len(fields) > 0:
            print("\n✅ TEST PASSED")
        else:
            print("\n⚠️ HTML generation incomplete")
            
    finally:
        if os.path.exists(test_file):
            os.remove(test_file)


def run_all_tests():
    """Run all tests"""
    print("\n" + "="*60)
    print("FORM REPLACEMENT SERVICE - TEST SUITE")
    print("="*60)
    
    try:
        test_dot_line_detection()
        test_field_extraction()
        test_dot_line_replacement()
        
        print("\n" + "="*60)
        print("✅ ALL TESTS COMPLETED")
        print("="*60 + "\n")
        
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()


if __name__ == '__main__':
    run_all_tests()
