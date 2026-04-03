"""
API Routes cho Form Replacement Service
Nâng cấp: Sử dụng Intelligent Detector để phát hiện tiêu đề, label, và tổ chức theo section
"""

from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Query
from sqlalchemy.orm import Session
from typing import List, Dict
import os
import json
from datetime import datetime

from app.db.session import get_db
from app.db.models import WordTemplate, User
from app.services.form_replacement import (
    DotLineDetector, 
    HTMLFieldReplacer,
    IntelligentDetector,
    SmartFormRenderer,
    DocumentStructurePreserver,
    FormLayoutRenderer,
    FormPageRenderer
)
from app.core.logger import logger
from app.core.auth import get_current_user
from app.core.file_utils import extract_clean_filename
from docx import Document

router = APIRouter(prefix="/api/form-replacement", tags=["form_replacement"])

# Upload directory
UPLOAD_DIR = "uploads"
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)


def resolve_effective_user_id(current_user: User, requested_user_id: int = None) -> int:
    """Resolve user scope from session user, allowing admin-only override."""
    if requested_user_id is None:
        return current_user.id

    if requested_user_id != current_user.id and not current_user.is_admin:
        raise HTTPException(status_code=403, detail="Không có quyền truy cập dữ liệu của người dùng khác")

    return requested_user_id


@router.post("/upload-with-intelligent-detection")
async def upload_form_with_intelligent_detection(
    file: UploadFile = File(...),
    user_id: int = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Upload file và detect form structure thông minh:
    - Phát hiện tiêu đề (headers)
    - Phát hiện label (fields)
    - Tổ chức theo section
    - Giữ nguyên cấu trúc document
    """
    
    # Kiểm tra định dạng file
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext != '.docx':
        raise HTTPException(
            status_code=400,
            detail=f"Hiện tại chỉ hỗ trợ file .docx. Bạn upload: {file_ext}"
        )
    
    try:
        effective_user_id = resolve_effective_user_id(current_user, user_id)
        
        # Lưu file
        file_path = os.path.join(UPLOAD_DIR, f"{effective_user_id}_{datetime.now().timestamp()}_{file.filename}")
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        # Intelligent detection
        doc = Document(file_path)
        parsed_form = IntelligentDetector.parse_document(doc)
        fields = IntelligentDetector.extract_field_list(parsed_form)
        
        if not parsed_form.fields:
            raise HTTPException(
                status_code=400,
                detail="Không tìm thấy bất kỳ trường nào. "
                       "Vui lòng chắc chắn file có chứa các ô trống dạng: '......', '____', '---'"
            )
        
        # Save template metadata
        template_metadata = {
            "sections": [
                {
                    "title": s.title,
                    "is_title": s.is_title,
                    "level": s.level,
                    "items_count": len(s.items)
                } for s in parsed_form.sections
            ],
            "fields": fields,
            "total_sections": len(parsed_form.sections),
            "total_fields": len(parsed_form.fields)
        }
        
        # Save template
        template = WordTemplate(
            user_id=effective_user_id,
            template_name=os.path.splitext(file.filename)[0],
            file_path=file_path,
            original_filename=file.filename,
            fields_json=json.dumps(fields)
        )
        db.add(template)
        db.commit()
        db.refresh(template)
        
        return {
            "status": "success",
            "template_id": template.id,
            "template_name": template.template_name,
            "file_type": file_ext,
            "fields_count": len(parsed_form.fields),
            "sections_count": len(parsed_form.sections),
            "fields": fields,
            "sections": template_metadata["sections"],
            "message": f"Upload thành công! Phát hiện {len(parsed_form.sections)} sections và {len(parsed_form.fields)} trường"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[Form Upload] Lỗi: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi: {str(e)}")


@router.get("/template/{template_id}/render-form-structured")
async def render_form_structured(
    template_id: int,
    user_id: int = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Render form HTML với cấu trúc được tổ chức theo section
    Giữ nguyên layout của document gốc
    """
    
    try:
        effective_user_id = resolve_effective_user_id(current_user, user_id)

        # Get template
        template = db.query(WordTemplate).filter(
            WordTemplate.id == template_id,
            WordTemplate.user_id == effective_user_id
        ).first()
        
        if not template:
            raise HTTPException(status_code=404, detail="Template không tồn tại")
        
        # Load and parse document
        doc = Document(template.file_path)
        parsed_form = IntelligentDetector.parse_document(doc)
        
        # Render form
        html_form = SmartFormRenderer.render_form_html(parsed_form)
        
        fields = IntelligentDetector.extract_field_list(parsed_form)
        
        return {
            "status": "success",
            "template_id": template_id,
            "template_name": template.template_name,
            "fields_count": len(parsed_form.fields),
            "sections_count": len(parsed_form.sections),
            "fields": fields,
            "html_form": html_form,
            "render_type": "structured",
            "message": "Form render thành công (structured layout)"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[Form Render] Lỗi: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi: {str(e)}")


@router.get("/template/{template_id}/render-form-inline")
async def render_form_inline(
    template_id: int,
    user_id: int = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Render form giữ nguyên cấu trúc gốc của document
    Thay thế placeholder trực tiếp trong text (inline replacement)
    """
    
    try:
        effective_user_id = resolve_effective_user_id(current_user, user_id)

        # Get template
        template = db.query(WordTemplate).filter(
            WordTemplate.id == template_id,
            WordTemplate.user_id == effective_user_id
        ).first()
        
        if not template:
            raise HTTPException(status_code=404, detail="Template không tồn tại")
        
        # Load and parse document
        doc = Document(template.file_path)
        parsed_form = IntelligentDetector.parse_document(doc)
        
        # Render with inline replacement
        html_form = SmartFormRenderer.render_form_with_inline_replacement(
            parsed_form.raw_content,
            parsed_form.fields
        )
        
        fields = IntelligentDetector.extract_field_list(parsed_form)
        
        return {
            "status": "success",
            "template_id": template_id,
            "template_name": template.template_name,
            "fields_count": len(parsed_form.fields),
            "fields": fields,
            "html_form": html_form,
            "render_type": "inline",
            "message": "Form render thành công (inline layout)"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[Form Render Inline] Lỗi: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi: {str(e)}")


@router.get("/template/{template_id}/render-form-document")
async def render_form_document(
    template_id: int,
    user_id: int = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Render form theo cấu trúc document - giống hình ảnh:
    - Tiêu đề: Non-editable text
    - Trường: Input boxes với borders
    - Layout: Chính xác như document gốc
    """
    
    try:
        effective_user_id = resolve_effective_user_id(current_user, user_id)

        # Get template
        template = db.query(WordTemplate).filter(
            WordTemplate.id == template_id,
            WordTemplate.user_id == effective_user_id
        ).first()
        
        if not template:
            raise HTTPException(status_code=404, detail="Template không tồn tại")
        
        # Load and parse document
        doc = Document(template.file_path)
        parsed_form = IntelligentDetector.parse_document(doc)
        
        # Render as document layout
        html_form = FormLayoutRenderer.render_form_document_style(parsed_form)
        
        fields = IntelligentDetector.extract_field_list(parsed_form)
        
        return {
            "status": "success",
            "template_id": template_id,
            "template_name": template.template_name,
            "fields_count": len(parsed_form.fields),
            "sections_count": len(parsed_form.sections),
            "fields": fields,
            "html_form": html_form,
            "render_type": "document",
            "message": "Form render thành công (document layout)"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[Form Render Document] Lỗi: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi: {str(e)}")


@router.get("/template/{template_id}/render-form-page")
async def render_form_page(
    template_id: int,
    user_id: int = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Render form as complete HTML page - ready to display
    Includes CSS, JS, buttons, etc.
    """
    
    try:
        effective_user_id = resolve_effective_user_id(current_user, user_id)

        # Get template
        template = db.query(WordTemplate).filter(
            WordTemplate.id == template_id,
            WordTemplate.user_id == effective_user_id
        ).first()
        
        if not template:
            raise HTTPException(status_code=404, detail="Template không tồn tại")
        
        # Load and parse document
        doc = Document(template.file_path)
        parsed_form = IntelligentDetector.parse_document(doc)
        
        # Render as complete page
        html_page = FormPageRenderer.render_complete_form_page(parsed_form)
        
        fields = IntelligentDetector.extract_field_list(parsed_form)
        
        return {
            "status": "success",
            "template_id": template_id,
            "template_name": template.template_name,
            "fields_count": len(parsed_form.fields),
            "sections_count": len(parsed_form.sections),
            "fields": fields,
            "html_page": html_page,
            "render_type": "page",
            "message": "Form page render thành công"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[Form Render Page] Lỗi: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi: {str(e)}")


@router.post("/upload-with-dotlines")
async def upload_form_with_dotlines(
    file: UploadFile = File(...),
    user_id: int = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Upload file và detect fields bằng dot-lines (legacy API)
    
    Thay vào đó, sử dụng /upload-with-intelligent-detection cho tính năng mới
    """
    
    # Kiểm tra định dạng file
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext != '.docx':
        raise HTTPException(
            status_code=400,
            detail=f"Hiện tại chỉ hỗ trợ file .docx. Bạn upload: {file_ext}"
        )
    
    try:
        effective_user_id = resolve_effective_user_id(current_user, user_id)
        
        # Lưu file
        file_path = os.path.join(UPLOAD_DIR, f"{effective_user_id}_{datetime.now().timestamp()}_{file.filename}")
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        # Detect dot-line placeholders
        doc = Document(file_path)
        
        placeholders = DotLineDetector.detect_from_document(doc)
        fields = DotLineDetector.extract_fields(placeholders)
        
        if not placeholders:
            raise HTTPException(
                status_code=400,
                detail="Không tìm thấy bất kỳ dot-lines (dòng chấm) nào trong file. "
                       "Vui lòng chắc chắn file có chứa các ô trống dạng: '......', '____', '---', vv"
            )
        
        # Save template
        template = WordTemplate(
            user_id=effective_user_id,
            template_name=os.path.splitext(file.filename)[0],
            file_path=file_path,
            original_filename=file.filename,
            fields_json=json.dumps([f['placeholder'] for f in fields])
        )
        db.add(template)
        db.commit()
        db.refresh(template)
        
        return {
            "status": "success",
            "template_id": template.id,
            "template_name": template.template_name,
            "file_type": file_ext,
            "fields_count": len(fields),
            "fields": fields,
            "placeholders_count": len(placeholders),
            "message": f"Upload thành công! Tìm thấy {len(fields)} trường từ dot-lines (legacy API)"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Lỗi upload file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi: {str(e)}")


@router.get("/template/{template_id}/render-form")
async def render_form_with_dotlines(
    template_id: int,
    user_id: int = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Render form HTML bằng cách replace dot-lines với input fields (legacy API)
    """
    
    try:
        effective_user_id = resolve_effective_user_id(current_user, user_id)

        # Get template
        template = db.query(WordTemplate).filter(
            WordTemplate.id == template_id,
            WordTemplate.user_id == effective_user_id
        ).first()
        
        if not template:
            raise HTTPException(status_code=404, detail="Template không tồn tại")
        
        # Load document
        doc = Document(template.file_path)
        
        # Detect placeholders
        placeholders = DotLineDetector.detect_from_document(doc)
        fields = DotLineDetector.extract_fields(placeholders)
        
        # Get paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Render form with replacements
        html_form = HTMLFieldReplacer.render_form_with_replacements(
            paragraphs,
            placeholders,
            fields
        )
        
        return {
            "status": "success",
            "template_id": template_id,
            "template_name": template.template_name,
            "fields_count": len(fields),
            "fields": fields,
            "html_form": html_form,
            "message": "Form render thành công (legacy API)"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Lỗi render form: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi: {str(e)}")


@router.get("/templates-with-dotlines")
async def get_dotline_templates(
    user_id: int = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Lấy danh sách templates - chỉ hỗ trợ dot-line templates"""

    effective_user_id = resolve_effective_user_id(current_user, user_id)
    
    templates = db.query(WordTemplate).filter(
        WordTemplate.user_id == effective_user_id
    ).all()
    
    return {
        "templates": [
            {
                "id": t.id,
                "name": t.template_name,
                "filename": extract_clean_filename(t.original_filename),
                "created_at": t.created_at.isoformat() if t.created_at else None
            }
            for t in templates
        ],
        "count": len(templates)
    }


@router.post("/submit-dotline-form")
async def submit_dotline_form(
    template_id: int = Query(...),
    user_id: int = Query(None),
    form_data: Dict = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Lưu submission từ form được render"""
    
    if not form_data:
        raise HTTPException(status_code=400, detail="Không có form data")

    _ = resolve_effective_user_id(current_user, user_id)
    
    # Có thể implement logic lưu form submission tại đây
    return {
        "status": "success",
        "message": "Form submitted successfully"
    }


@router.get("/templates-with-dotlines")
async def get_dotline_templates(
    user_id: int = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Lấy danh sách templates được detect bằng dot-lines"""

    effective_user_id = resolve_effective_user_id(current_user, user_id)
    
    templates = db.query(WordTemplate).filter(
        WordTemplate.user_id == effective_user_id
    ).all()
    
    return {
        "templates": [
            {
                "id": t.id,
                "name": t.template_name,
                "filename": extract_clean_filename(t.original_filename),
                "fields_count": len(json.loads(t.fields_json)) if t.fields_json else 0,
                "created_at": t.created_at.isoformat(),
                "upload_method": "dot-lines"  # Mark as dot-line based
            }
            for t in templates
        ]
    }


@router.post("/submit-dotline-form")
async def submit_dotline_form(
    form_data: Dict,
    template_id: int = Query(...),
    user_id: int = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Submit form data từ dot-line form"""
    
    try:
        effective_user_id = resolve_effective_user_id(current_user, user_id)

        # Similar to existing submit logic 
        # Just with dot-line specific handling
        
        from app.db.models import WordSubmission, Entry, Field, Form
        
        # Get template
        template = db.query(WordTemplate).filter(
            WordTemplate.id == template_id
        ).first()
        
        if not template:
            raise HTTPException(status_code=404, detail="Template không tồn tại")

        if template.user_id != effective_user_id and not current_user.is_admin:
            raise HTTPException(status_code=403, detail="Không có quyền sử dụng template này")
        
        # Get or create form
        form = db.query(Form).filter(Form.id == 1).first()
        if not form:
            form = Form(id=1, user_id=effective_user_id, form_name="Dot-Line Form", description="Form from dot-lines")
            db.add(form)
            db.commit()
            db.refresh(form)
        
        # Create submission
        submission = WordSubmission(
            template_id=template_id,
            user_id=effective_user_id,
            form_id=form.id,
            submission_data=json.dumps(form_data)
        )
        db.add(submission)
        db.commit()
        db.refresh(submission)
        
        return {
            "status": "success",
            "submission_id": submission.id,
            "message": "Submit thành công"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Lỗi submit form: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi: {str(e)}")
