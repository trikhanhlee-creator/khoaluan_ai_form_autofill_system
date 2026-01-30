"""
Word Document Parser - Trích xuất fields từ file Word
"""

from docx import Document
from typing import List, Dict
import re
from datetime import datetime
import os


class WordField:
    """Model cho một field từ file Word"""
    def __init__(self, name: str, field_type: str = "text", label: str = "", order: int = 0):
        self.name = name
        self.field_type = field_type
        self.label = label
        self.order = order

    def to_dict(self):
        return {
            "name": self.name,
            "field_type": self.field_type,
            "label": self.label,
            "order": self.order
        }


class WordParser:
    """Parser để trích xuất fields từ file Word (.docx)"""
    
    FIELD_TYPE_KEYWORDS = {
        'năm': 'number',
        'ngày': 'date',
        'điện thoại': 'phone',
        'email': 'email',
        'địa chỉ': 'text',
        'tên': 'text',
        'họ': 'text',
        'số': 'number',
        'năm sinh': 'number',
    }
    
    def __init__(self, file_path: str):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File không tồn tại: {file_path}")
        
        if not file_path.lower().endswith('.docx'):
            raise ValueError("Chỉ hỗ trợ file .docx")
        
        self.file_path = file_path
        self.doc = Document(file_path)
        self.fields: List[WordField] = []
    
    def detect_field_type(self, text: str) -> str:
        """Phát hiện kiểu dữ liệu từ nhãn trường"""
        text_lower = text.lower()
        
        for keyword, field_type in self.FIELD_TYPE_KEYWORDS.items():
            if keyword in text_lower:
                return field_type
        
        # Mặc định
        return "text"
    
    def parse_paragraphs(self) -> List[WordField]:
        """Trích xuất fields từ paragraphs (định dạng: "Nhãn: " hoặc "Nhãn:")"""
        fields = []
        order = 0
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Xóa dấu chấm câu ở cuối
            label = text.rstrip('.:')
            
            # Nếu text kết thúc bằng ":" hoặc chứa ":", đó là một field
            if ':' in text or text.endswith(':'):
                # Tạo tên field từ label (xóa dấu, chuyển thành snake_case)
                field_name = label.lower()
                field_name = re.sub(r'[^\w\s]', '', field_name)  # Xóa dấu tiếng Việt
                field_name = field_name.strip()
                field_name = '_'.join(field_name.split())  # Chuyển thành snake_case
                
                # Phát hiện kiểu
                field_type = self.detect_field_type(label)
                
                # Tạo field
                field = WordField(
                    name=field_name,
                    field_type=field_type,
                    label=label,
                    order=order
                )
                fields.append(field)
                order += 1
        
        return fields
    
    def parse_tables(self) -> List[WordField]:
        """Trích xuất fields từ bảng (cột 1 = nhãn, cột 2 = giá trị)"""
        fields = []
        order = 0
        
        for table in self.doc.tables:
            # Nếu bảng có 2 cột (nhãn, giá trị)
            if len(table.columns) >= 2:
                for row in table.rows:
                    if len(row.cells) >= 2:
                        label = row.cells[0].text.strip()
                        
                        if label:
                            # Tạo tên field
                            field_name = label.lower()
                            field_name = re.sub(r'[^\w\s]', '', field_name)
                            field_name = field_name.strip()
                            field_name = '_'.join(field_name.split())
                            
                            # Phát hiện kiểu
                            field_type = self.detect_field_type(label)
                            
                            # Tạo field
                            field = WordField(
                                name=field_name,
                                field_type=field_type,
                                label=label,
                                order=order
                            )
                            fields.append(field)
                            order += 1
        
        return fields
    
    def parse(self) -> List[WordField]:
        """Parse file Word và trích xuất tất cả fields"""
        # Ưu tiên paragraphs, nếu không có thì dùng tables
        fields = self.parse_paragraphs()
        
        if not fields:
            fields = self.parse_tables()
        
        self.fields = fields
        return fields
    
    def to_dict_list(self) -> List[Dict]:
        """Chuyển fields thành list dict"""
        return [field.to_dict() for field in self.fields]
    
    def get_metadata(self) -> Dict:
        """Lấy metadata của file Word"""
        core_props = self.doc.core_properties
        return {
            "title": core_props.title or os.path.basename(self.file_path),
            "author": core_props.author or "Unknown",
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None,
            "paragraphs_count": len(self.doc.paragraphs),
            "tables_count": len(self.doc.tables),
            "fields_count": len(self.fields),
        }


# Test
if __name__ == "__main__":
    parser = WordParser('C:\\Users\\KHANH\\Downloads\\testform.docx')
    fields = parser.parse()
    
    print("=== PARSED FIELDS ===")
    for field in fields:
        print(f"- {field.label} ({field.name}): {field.field_type}")
    
    print("\n=== METADATA ===")
    print(parser.get_metadata())
