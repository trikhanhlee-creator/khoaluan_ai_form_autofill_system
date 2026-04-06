"""
Multi-format File Parser - Hỗ trợ parse các định dạng file khác nhau (.docx, .pdf, .xlsx, .xls, .csv, .txt)
"""

from abc import ABC, abstractmethod
from typing import List, Dict
import re
import os
import json
from datetime import datetime


class FileField:
    """Model cho một field được trích xuất từ file"""
    def __init__(self, name: str, field_type: str = "text", label: str = "", order: int = 0):
        self.name = name
        self.field_type = field_type
        self.label = label
        self.order = order

    def to_dict(self):
        return {
            "name": self.name,
            "field_type": self.field_type,
            "label": self.label,
            "order": self.order
        }


class BaseFileParser(ABC):
    """Base class cho tất cả file parsers"""
    
    FIELD_TYPE_KEYWORDS = {
        'ngày': 'date',
        'năm sinh': 'number',
        'năm': 'number',
        'điện thoại': 'phone',
        'email': 'email',
        'địa chỉ': 'text',
        'tên': 'text',
        'họ': 'text',
        'số': 'number',
    }
    
    SUPPORTED_EXTENSIONS = []
    
    def __init__(self, file_path: str):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File không tồn tại: {file_path}")
        
        self.file_path = file_path
        self.fields: List[FileField] = []
    
    def clean_field_label(self, text: str) -> str:
        """Clean field label by removing separator characters at end"""
        text = text.strip()
        
        if not text:
            return text
        
        # Remove wrapping and embedded separator patterns
        text = re.sub(r'\s*[\(\[\{]\s*[\.─\-_*]+\s*[\)\]\}]\s*$', '', text)
        
        # Remove trailing separators/delimiters and Unicode smart quotes
        text = re.sub(r'[\s.:\,;!)\]\}»"\'─\-_*~`(\u201c\u201d\u2018\u2019]+$', '', text)
        
        # Remove leading special characters
        text = re.sub(r'^[\s«\(\[\{\'"─\-_*~`(\u201c\u201d\u2018\u2019]+', '', text)
        
        text = text.strip()
        return text
    
    def detect_field_type(self, text: str) -> str:
        """Phát hiện kiểu dữ liệu từ nhãn trường"""
        text_lower = text.lower()
        
        for keyword, field_type in self.FIELD_TYPE_KEYWORDS.items():
            if keyword in text_lower:
                return field_type
        
        return "text"
    
    def create_field_name(self, label: str) -> str:
        """Tạo field name từ label (snake_case)"""
        field_name = label.lower()
        field_name = re.sub(r'[^\w\s]', '', field_name)
        field_name = field_name.strip()
        field_name = '_'.join(field_name.split())
        return field_name
    
    @abstractmethod
    def parse(self) -> List[FileField]:
        """Parse file và trích xuất fields"""
        pass
    
    def get_metadata(self) -> Dict:
        """Lấy metadata của file"""
        return {
            "title": os.path.basename(self.file_path),
            "file_type": os.path.splitext(self.file_path)[1].lower(),
            "fields_count": len(self.fields),
        }


class DocxParser(BaseFileParser):
    """Parser cho file Word (.docx)"""
    
    SUPPORTED_EXTENSIONS = ['.docx']
    
    def __init__(self, file_path: str):
        super().__init__(file_path)
        from docx import Document
        self.doc = Document(file_path)
        self.detected_title = ""

    PLACEHOLDER_RE = re.compile(r'(\.{3,}|_{3,}|…{2,}|-{3,}|─{3,}|‒{3,}|–{3,}|—{3,}|\u2026{2,})')
    CHECKBOX_RE = re.compile(r'(?:\[\s*\]|\(\s*\)|□|☐|☑|✓)')
    DATE_TEMPLATE_RE = re.compile(r'ngày\s*[.\-_/…_]{0,10}\s*tháng\s*[.\-_/…_]{0,10}\s*năm', re.IGNORECASE)

    HEADING_PREFIXES = (
        'cộng hòa xã hội chủ nghĩa',
        'độc lập - tự do - hạnh phúc',
        'đơn ',
        'mẫu số',
    )

    NON_FIELD_PREFIXES = (
        'cộng hòa xã hội chủ nghĩa',
        'độc lập - tự do - hạnh phúc',
        'kính gửi',
        'tôi xin chân thành cảm ơn',
        'người viết đơn',
        'ký tên',
        'xác nhận của',
        'nơi nhận',
        'ghi chú',
        'hướng dẫn',
        'thông qua',
    )

    def _normalize_line(self, text: str) -> str:
        return re.sub(r'\s+', ' ', (text or '').strip())

    def _has_placeholder_hint(self, text: str) -> bool:
        normalized = self._normalize_line(text)
        if not normalized:
            return False
        if self.PLACEHOLDER_RE.search(normalized):
            return True
        if self.CHECKBOX_RE.search(normalized):
            return True
        if self.DATE_TEMPLATE_RE.search(normalized):
            return True
        if re.search(r'[:：]\s*$', normalized):
            return True
        if re.search(r'[:：]\s*[^\s]{0,2}$', normalized):
            return True
        return False

    def _is_likely_heading(self, para, text: str) -> bool:
        normalized = self._normalize_line(text)
        if not normalized:
            return False

        lower = normalized.lower()
        words = normalized.split()
        is_center = getattr(para, 'alignment', None) == 1
        style_name = (getattr(getattr(para, 'style', None), 'name', '') or '').lower()
        is_heading_style = 'heading' in style_name or 'title' in style_name
        is_upper_like = normalized == normalized.upper() and len(words) <= 14

        return (
            len(words) <= 18
            and len(normalized) <= 120
            and not self._has_placeholder_hint(normalized)
            and (
                is_center
                or is_heading_style
                or is_upper_like
                or lower.startswith(self.HEADING_PREFIXES)
            )
        )

    def _is_non_field_content(self, text: str, label: str) -> bool:
        normalized = self._normalize_line(text)
        if not normalized:
            return True

        if self._is_signature_date_placeholder_line(normalized):
            return True

        lower = normalized.lower()
        words = normalized.split()

        if lower.startswith(self.NON_FIELD_PREFIXES) and not self._has_placeholder_hint(normalized):
            return True

        if len(label) > 120:
            return True

        if not self._has_placeholder_hint(normalized):
            if len(words) > 18:
                return True
            if normalized.count(',') >= 2:
                return True
            if re.search(r'[\.!?;]\s+\S', normalized) and len(words) > 12:
                return True

        return False

    def _looks_like_paragraph_field(self, text: str) -> bool:
        normalized = self._normalize_line(text)
        if not normalized:
            return False

        if self._has_placeholder_hint(normalized):
            return True

        words = normalized.split()
        if len(words) <= 8 and len(normalized) <= 60 and re.search(r'[:：]', normalized):
            return True

        return False

    def _is_signature_date_placeholder_line(self, text: str) -> bool:
        normalized = self._normalize_line(text)
        if not normalized or not self.DATE_TEMPLATE_RE.search(normalized):
            return False

        lower = normalized.lower()

        # Explicitly labeled date fields are likely real inputs.
        if ':' in normalized or '：' in normalized:
            return False
        if any(keyword in lower for keyword in ('ngày nộp', 'nộp đơn', 'ngày lập', 'ngày ký')):
            return False

        compact = re.sub(r'[.\-_/…_‒–—,;:]', ' ', normalized)
        words = re.findall(r'[a-zà-ỹ]+', compact.lower())
        core_words = [w for w in words if w not in {'ngày', 'tháng', 'năm'}]

        if not words or not core_words:
            return True

        # Typical signature footer: optional location before comma + ngày/tháng/năm template.
        if ',' in normalized and len(core_words) <= 2:
            return True

        if re.match(r'^[\W_]+', normalized) and len(core_words) <= 3:
            return True

        return False

    def _create_field(self, label: str, order: int) -> FileField | None:
        cleaned = self.clean_field_label(label)
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        if not cleaned or len(cleaned) < 2:
            return None

        field_name = self.create_field_name(cleaned)
        if not field_name:
            return None

        return FileField(
            name=field_name,
            field_type=self.detect_field_type(cleaned),
            label=cleaned,
            order=order,
        )

    def _normalize_placeholder_label(self, label: str) -> str:
        value = re.sub(r'\s+', ' ', (label or '').strip())
        value = re.sub(r'^(và|hoặc|là|là:|tại|ở|cho|của|theo|về)\s+', '', value, flags=re.IGNORECASE)
        value = value.strip(' .,:;')

        words = value.split()
        if len(words) > 8:
            value = ' '.join(words[-6:])

        lower = value.lower()
        if any(phrase in lower for phrase in ('tôi tên là', 'họ tên', 'họ và tên')):
            return 'Họ và tên'
        if 'sinh ngày' in lower or 'ngày sinh' in lower:
            return 'Ngày sinh'
        if 'chỗ ở' in lower or 'địa chỉ' in lower:
            return 'Địa chỉ hiện tại'
        if 'điện thoại' in lower:
            return 'Số điện thoại liên hệ'
        if 'vị trí' in lower and ('ứng tuyển' in lower or 'tuyển' in lower or len(lower.split()) <= 5):
            return 'Vị trí ứng tuyển'
        if 'công ty' in lower:
            return 'Công ty ứng tuyển'
        if 'tốt nghiệp' in lower and 'loại' in lower:
            return 'Xếp loại tốt nghiệp'
        if re.search(r'(?:^|\s)(tại\s+)?trường$', lower):
            return 'Trường tốt nghiệp'
        if 'khóa học' in lower:
            return 'Khóa học đã tham gia'
        if self.DATE_TEMPLATE_RE.search(lower) or ('ngày' in lower and 'tháng' in lower and 'năm' in lower):
            return 'Ngày nộp đơn'

        return value

    def _extract_label_from_placeholder_context(self, prefix: str, suffix: str) -> str:
        candidate = self._normalize_line(prefix)

        if ':' in candidate or '：' in candidate:
            parts = re.split(r'[:：]', candidate)
            after_colon = parts[-1].strip() if parts else ''
            before_colon = parts[-2].strip() if len(parts) >= 2 else ''
            candidate = after_colon or before_colon

        candidate = re.split(r'[.!?;]', candidate)[-1].strip()

        if not candidate:
            candidate = re.split(r'[,:;.!?]', self._normalize_line(suffix))[0].strip()

        return self._normalize_placeholder_label(candidate)

    def _extract_fields_from_placeholder_line(self, text: str, order_start: int = 0) -> List[FileField]:
        normalized = self._normalize_line(text)
        if not normalized or not self._has_placeholder_hint(normalized):
            return []

        if self._is_signature_date_placeholder_line(normalized):
            return []

        fields: List[FileField] = []
        seen_names: set[str] = set()

        checkbox_match = self.CHECKBOX_RE.search(normalized)
        if checkbox_match:
            checkbox_label = self._extract_label_from_placeholder_context(
                normalized[:checkbox_match.start()],
                normalized[checkbox_match.end():],
            )
            checkbox_field = self._create_field(checkbox_label, order_start + len(fields))
            if checkbox_field and checkbox_field.name not in seen_names:
                seen_names.add(checkbox_field.name)
                fields.append(checkbox_field)

        matches = list(self.PLACEHOLDER_RE.finditer(normalized))
        for idx, match in enumerate(matches):
            prev_end = matches[idx - 1].end() if idx > 0 else 0
            next_start = matches[idx + 1].start() if idx + 1 < len(matches) else len(normalized)

            prefix = normalized[prev_end:match.start()].strip() if idx > 0 else normalized[:match.start()].strip()
            suffix = normalized[match.end():next_start].strip()

            label = self._extract_label_from_placeholder_context(prefix, suffix)
            if not label and idx == 0:
                label = self._extract_label_from_placeholder_context(
                    normalized[:match.start()],
                    normalized[match.end():],
                )

            field = self._create_field(label, order_start + len(fields)) if label else None
            if field and field.name not in seen_names:
                seen_names.add(field.name)
                fields.append(field)

        if not fields and self.DATE_TEMPLATE_RE.search(normalized) and not self._is_signature_date_placeholder_line(normalized):
            date_field = self._create_field('Ngày nộp đơn', order_start)
            if date_field:
                fields.append(date_field)

        return fields

    def _dedupe_fields(self, fields: List[FileField]) -> List[FileField]:
        unique_fields: List[FileField] = []
        seen: set[str] = set()

        for field in fields:
            key = field.name.strip().lower()
            if not key or key in seen:
                continue
            seen.add(key)
            unique_fields.append(field)

        for idx, field in enumerate(unique_fields):
            field.order = idx

        return unique_fields

    def _detect_document_title(self) -> str:
        title_candidates: List[str] = []

        for para in self.doc.paragraphs[:20]:
            text = self._normalize_line(para.text)
            if not text:
                continue

            if self._is_likely_heading(para, text):
                title_candidates.append(text)
                if text.lower().startswith('đơn ') or len(title_candidates) >= 2:
                    break

        if not title_candidates:
            return ""

        doc_title = next((t for t in title_candidates if t.lower().startswith('đơn ')), None)
        if doc_title:
            return doc_title

        if len(title_candidates) > 1:
            return ' - '.join(title_candidates[:2])

        return title_candidates[0]
    
    def parse_paragraphs(self) -> List[FileField]:
        """Trích xuất fields từ paragraphs"""
        fields = []
        self.detected_title = self._detect_document_title()
        
        for para in self.doc.paragraphs:
            text = self._normalize_line(para.text)
            if not text:
                continue

            if self._is_likely_heading(para, text):
                continue

            if self._has_placeholder_hint(text):
                extracted = self._extract_fields_from_placeholder_line(text, len(fields))
                if extracted:
                    fields.extend(extracted)
                    continue

            if not self._looks_like_paragraph_field(text):
                continue

            label = self.clean_field_label(text)
            if not label or len(label) < 2:
                continue

            if self._is_non_field_content(text, label):
                continue

            field = self._create_field(label, len(fields))
            if field:
                fields.append(field)

        return self._dedupe_fields(fields)
    
    def parse_tables(self) -> List[FileField]:
        """Trích xuất fields từ bảng"""
        fields = []
 
        for table in self.doc.tables:
            for row in table.rows:
                cells = row.cells
                for cell_idx, cell in enumerate(cells):
                    text = self._normalize_line(cell.text)
                    if not text:
                        continue

                    if self._has_placeholder_hint(text):
                        extracted = self._extract_fields_from_placeholder_line(text, len(fields))
                        if extracted:
                            fields.extend(extracted)
                        elif cell_idx > 0:
                            left_text = self._normalize_line(cells[cell_idx - 1].text)
                            if left_text and not self._has_placeholder_hint(left_text):
                                label = self.clean_field_label(left_text)
                                if label and not self._is_non_field_content(left_text, label):
                                    left_field = self._create_field(label, len(fields))
                                    if left_field:
                                        fields.append(left_field)

                    if cell_idx + 1 < len(cells):
                        right_text = self._normalize_line(cells[cell_idx + 1].text)
                        if right_text and self._has_placeholder_hint(right_text) and not self._has_placeholder_hint(text):
                            label = self.clean_field_label(text)
                            if label and not self._is_non_field_content(text, label):
                                right_field = self._create_field(label, len(fields))
                                if right_field:
                                    fields.append(right_field)

        return self._dedupe_fields(fields)
    
    def parse_all_text_content(self) -> List[FileField]:
        """Fallback: Trích xuất tất cả text content thành fields"""
        fields = []
        order = 0
        
        for para in self.doc.paragraphs:
            text = self._normalize_line(para.text)
            if not text or len(text) < 2 or len(text) > 500:
                continue

            if self._is_likely_heading(para, text):
                continue
            
            label = self.clean_field_label(text)
            if not label or len(label) < 2:
                continue

            if self._is_non_field_content(text, label):
                continue
            
            field_name = self.create_field_name(label)
            if not field_name:
                continue
            
            field_type = self.detect_field_type(label)
            field = FileField(
                name=field_name,
                field_type=field_type,
                label=label,
                order=order
            )
            fields.append(field)
            order += 1
            
            if order >= 20:
                break

        return self._dedupe_fields(fields)
    
    def parse(self) -> List[FileField]:
        """Parse file Word"""
        paragraph_fields = self.parse_paragraphs()
        table_fields = self.parse_tables()

        fields = self._dedupe_fields(paragraph_fields + table_fields)
        if not fields:
            fields = self.parse_all_text_content()

        self.fields = self._dedupe_fields(fields)
        return self.fields
    
    def get_metadata(self) -> Dict:
        """Lấy metadata của file Word"""
        metadata = super().get_metadata()
        core_props = self.doc.core_properties
        title = core_props.title or self.detected_title or os.path.basename(self.file_path)
        metadata.update({
            "title": title,
            "author": core_props.author or "Unknown",
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None,
            "paragraphs_count": len(self.doc.paragraphs),
            "tables_count": len(self.doc.tables),
        })
        return metadata


class PdfParser(BaseFileParser):
    """Parser cho file PDF (.pdf)"""
    
    SUPPORTED_EXTENSIONS = ['.pdf']
    
    def __init__(self, file_path: str):
        super().__init__(file_path)
        import pdfplumber
        self.pdf = pdfplumber.open(file_path)
    
    def parse_text_lines(self) -> List[FileField]:
        """Trích xuất fields từ text trong PDF"""
        fields = []
        order = 0
        
        all_text = ""
        for page in self.pdf.pages:
            all_text += page.extract_text() or ""
        
        lines = all_text.split('\n')
        
        for line in lines:
            text = line.strip()
            if not text or len(text) > 500:
                continue
            
            has_separator = any(sep in text for sep in [':', '.', '─', '(', '[', '{', '"', "'", '*', '_', '-'])
            is_short_label = (2 <= len(text) < 50 and len(text.split()) <= 15)
            
            if not (has_separator or is_short_label):
                continue
            
            label = self.clean_field_label(text)
            if not label or len(label) < 2:
                continue
            
            field_name = self.create_field_name(label)
            if not field_name:
                continue
            
            field_type = self.detect_field_type(label)
            field = FileField(
                name=field_name,
                field_type=field_type,
                label=label,
                order=order
            )
            fields.append(field)
            order += 1
            
            if order >= 20:
                break
        
        return fields
    
    def parse(self) -> List[FileField]:
        """Parse file PDF"""
        try:
            fields = self.parse_text_lines()
            self.fields = fields
            return fields
        finally:
            self.pdf.close()
    
    def get_metadata(self) -> Dict:
        """Lấy metadata của file PDF"""
        metadata = super().get_metadata()
        metadata.update({
            "pages_count": len(self.pdf.pages),
        })
        return metadata


class XlsxParser(BaseFileParser):
    """Parser cho file Excel (.xlsx)"""
    
    SUPPORTED_EXTENSIONS = ['.xlsx', '.xls']
    
    def __init__(self, file_path: str):
        super().__init__(file_path)
        self.file_ext = os.path.splitext(file_path)[1].lower()
        self.workbook = None
        self.worksheet = None
        self.xlrd_workbook = None
        self.xlrd_sheet = None

        if self.file_ext == '.xlsx':
            from openpyxl import load_workbook
            self.workbook = load_workbook(file_path, data_only=True)
            self.worksheet = self.workbook.active
        elif self.file_ext == '.xls':
            import xlrd
            self.xlrd_workbook = xlrd.open_workbook(file_path)
            if self.xlrd_workbook.nsheets == 0:
                raise ValueError('File .xls không có sheet nào')
            self.xlrd_sheet = self.xlrd_workbook.sheet_by_index(0)
        else:
            raise ValueError(f'Định dạng Excel không được hỗ trợ: {self.file_ext}')

    def _parse_xlsx_sheet(self) -> List[FileField]:
        """Trích xuất fields từ sheet .xlsx"""
        fields = []
        order = 0

        header_row = self.worksheet[1]
        has_headers = False

        for cell in header_row:
            if cell.value:
                text = str(cell.value).strip()
                if text and len(text) > 0 and len(text) < 100:
                    has_headers = True
                    break

        if has_headers:
            for cell in header_row:
                if cell.value:
                    text = str(cell.value).strip()
                    if not text or len(text) < 1:
                        continue

                    label = self.clean_field_label(text)
                    if not label or len(label) < 1:
                        continue

                    field_name = self.create_field_name(label)
                    if not field_name:
                        continue

                    field_type = self.detect_field_type(label)
                    field = FileField(
                        name=field_name,
                        field_type=field_type,
                        label=label,
                        order=order
                    )
                    fields.append(field)
                    order += 1

        if not fields:
            for row in self.worksheet.iter_rows(min_row=1, max_row=20, values_only=True):
                if row and row[0]:
                    text = str(row[0]).strip()
                    if not text or len(text) > 500:
                        continue

                    has_separator = any(sep in text for sep in [':', '.', '─', '(', '[', '{'])
                    is_short_label = (2 <= len(text) < 50 and len(text.split()) <= 15)

                    if not (has_separator or is_short_label):
                        continue

                    label = self.clean_field_label(text)
                    if not label or len(label) < 2:
                        continue

                    field_name = self.create_field_name(label)
                    if not field_name:
                        continue

                    field_type = self.detect_field_type(label)
                    field = FileField(
                        name=field_name,
                        field_type=field_type,
                        label=label,
                        order=order
                    )
                    fields.append(field)
                    order += 1

                    if order >= 20:
                        break

        return fields

    def _parse_xls_sheet(self) -> List[FileField]:
        """Trích xuất fields từ sheet .xls"""
        fields = []
        order = 0

        if not self.xlrd_sheet or self.xlrd_sheet.nrows == 0:
            return fields

        header_row = self.xlrd_sheet.row_values(0)
        has_headers = any(str(cell).strip() and len(str(cell).strip()) < 100 for cell in header_row)

        if has_headers:
            for cell in header_row:
                text = str(cell).strip()
                if not text or len(text) < 1:
                    continue

                label = self.clean_field_label(text)
                if not label or len(label) < 1:
                    continue

                field_name = self.create_field_name(label)
                if not field_name:
                    continue

                field_type = self.detect_field_type(label)
                field = FileField(
                    name=field_name,
                    field_type=field_type,
                    label=label,
                    order=order
                )
                fields.append(field)
                order += 1

        if not fields:
            max_rows = min(20, self.xlrd_sheet.nrows)
            for row_idx in range(max_rows):
                row_values = self.xlrd_sheet.row_values(row_idx)
                if not row_values:
                    continue

                text = str(row_values[0]).strip()
                if not text or len(text) > 500:
                    continue

                has_separator = any(sep in text for sep in [':', '.', '─', '(', '[', '{'])
                is_short_label = (2 <= len(text) < 50 and len(text.split()) <= 15)

                if not (has_separator or is_short_label):
                    continue

                label = self.clean_field_label(text)
                if not label or len(label) < 2:
                    continue

                field_name = self.create_field_name(label)
                if not field_name:
                    continue

                field_type = self.detect_field_type(label)
                field = FileField(
                    name=field_name,
                    field_type=field_type,
                    label=label,
                    order=order
                )
                fields.append(field)
                order += 1

                if order >= 20:
                    break

        return fields
    
    def parse_sheets(self) -> List[FileField]:
        """Trích xuất fields từ dòng đầu tiên hoặc cột đầu tiên"""
        if self.file_ext == '.xlsx':
            return self._parse_xlsx_sheet()
        if self.file_ext == '.xls':
            return self._parse_xls_sheet()
        return []
    
    def parse(self) -> List[FileField]:
        """Parse file Excel"""
        try:
            fields = self.parse_sheets()
            self.fields = fields
            return fields
        finally:
            if self.workbook is not None:
                self.workbook.close()
    
    def get_metadata(self) -> Dict:
        """Lấy metadata của file Excel"""
        metadata = super().get_metadata()
        if self.file_ext == '.xlsx' and self.workbook is not None and self.worksheet is not None:
            metadata.update({
                "sheet_names": self.workbook.sheetnames,
                "active_sheet": self.worksheet.title,
                "max_row": self.worksheet.max_row,
                "max_column": self.worksheet.max_column,
            })
        elif self.file_ext == '.xls' and self.xlrd_workbook is not None and self.xlrd_sheet is not None:
            metadata.update({
                "sheet_names": self.xlrd_workbook.sheet_names(),
                "active_sheet": self.xlrd_sheet.name,
                "max_row": self.xlrd_sheet.nrows,
                "max_column": self.xlrd_sheet.ncols,
            })
        return metadata


class CsvParser(BaseFileParser):
    """Parser cho file CSV (.csv)"""
    
    SUPPORTED_EXTENSIONS = ['.csv']
    
    def __init__(self, file_path: str):
        super().__init__(file_path)
        self.rows = self._read_csv()
    
    def _read_csv(self):
        """Đọc file CSV"""
        import csv
        rows = []
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append(row)
        except:
            # Try with different encoding
            with open(self.file_path, 'r', encoding='latin-1') as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append(row)
        return rows
    
    def parse(self) -> List[FileField]:
        """Parse file CSV"""
        fields = []
        order = 0
        
        if not self.rows:
            return fields
        
        # Cách 1: Dùng header row
        header_row = self.rows[0]
        has_headers = all(cell and len(str(cell).strip()) > 0 for cell in header_row)
        
        if has_headers:
            for cell in header_row:
                if cell:
                    text = str(cell).strip()
                    if not text:
                        continue
                    
                    label = self.clean_field_label(text)
                    if not label or len(label) < 1:
                        continue
                    
                    field_name = self.create_field_name(label)
                    if not field_name:
                        continue
                    
                    field_type = self.detect_field_type(label)
                    field = FileField(
                        name=field_name,
                        field_type=field_type,
                        label=label,
                        order=order
                    )
                    fields.append(field)
                    order += 1
        
        # Cách 2: Nếu không có header, lấy từ cột đầu tiên
        if not fields:
            for row in self.rows[:20]:
                if row and row[0]:
                    text = str(row[0]).strip()
                    if not text or len(text) > 500:
                        continue
                    
                    has_separator = any(sep in text for sep in [':', '.', '─', '(', '[', '{'])
                    is_short_label = (2 <= len(text) < 50 and len(text.split()) <= 15)
                    
                    if not (has_separator or is_short_label):
                        continue
                    
                    label = self.clean_field_label(text)
                    if not label or len(label) < 2:
                        continue
                    
                    field_name = self.create_field_name(label)
                    if not field_name:
                        continue
                    
                    field_type = self.detect_field_type(label)
                    field = FileField(
                        name=field_name,
                        field_type=field_type,
                        label=label,
                        order=order
                    )
                    fields.append(field)
                    order += 1
                    
                    if order >= 20:
                        break
        
        self.fields = fields
        return fields


class TxtParser(BaseFileParser):
    """Parser cho file text tin (.txt)"""
    
    SUPPORTED_EXTENSIONS = ['.txt']
    
    def __init__(self, file_path: str):
        super().__init__(file_path)
        self.content = self._read_file()
    
    def _read_file(self):
        """Đọc file text"""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except:
            with open(self.file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def parse(self) -> List[FileField]:
        """Parse file text"""
        fields = []
        order = 0
        
        lines = self.content.split('\n')
        
        for line in lines:
            text = line.strip()
            if not text or len(text) > 500:
                continue
            
            has_separator = any(sep in text for sep in [':', '.', '─', '(', '[', '{', '"', "'", '*', '_', '-'])
            is_short_label = (2 <= len(text) < 50 and len(text.split()) <= 15)
            
            if not (has_separator or is_short_label):
                continue
            
            label = self.clean_field_label(text)
            if not label or len(label) < 2:
                continue
            
            field_name = self.create_field_name(label)
            if not field_name:
                continue
            
            field_type = self.detect_field_type(label)
            field = FileField(
                name=field_name,
                field_type=field_type,
                label=label,
                order=order
            )
            fields.append(field)
            order += 1
            
            if order >= 20:
                break
        
        self.fields = fields
        return fields


class FileParserFactory:
    """Factory class để tạo parser phù hợp dựa trên file extension"""
    
    PARSERS = {
        '.docx': DocxParser,
        '.pdf': PdfParser,
        '.xlsx': XlsxParser,
        '.xls': XlsxParser,
        '.csv': CsvParser,
        '.txt': TxtParser,
    }
    
    SUPPORTED_EXTENSIONS = list(PARSERS.keys())
    
    @classmethod
    def create_parser(cls, file_path: str) -> BaseFileParser:
        """Tạo parser phù hợp dựa trên file extension"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in cls.PARSERS:
            raise ValueError(f"Không hỗ trợ định dạng file: {file_ext}. Các định dạng được hỗ trợ: {', '.join(cls.SUPPORTED_EXTENSIONS)}")
        
        parser_class = cls.PARSERS[file_ext]
        return parser_class(file_path)
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """Lấy danh sách các extension được hỗ trợ"""
        return cls.SUPPORTED_EXTENSIONS
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Kiểm tra xem file có được hỗ trợ không"""
        file_ext = os.path.splitext(file_path)[1].lower()
        return file_ext in cls.PARSERS
